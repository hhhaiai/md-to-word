from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.oxml.shared import OxmlElement, qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml.parser import parse_xml
from typing import Dict, Any
import re

from config import DocumentConfig

class WordPostprocessor:
    """Word文档后处理器，对pandoc生成的Word文档应用GB/T 9704-2012格式要求"""
    
    def __init__(self):
        self.config = DocumentConfig()
        self.doc = None
    
    def apply_formatting(self, docx_path: str, metadata: Dict[str, Any], original_markdown: str = None) -> str:
        """
        对pandoc生成的Word文档应用公文格式
        
        Args:
            docx_path: pandoc生成的Word文档路径
            metadata: 包含标题、附件等元数据的字典
            original_markdown: 原始markdown内容，用于判断列表层级
            
        Returns:
            处理后的Word文档路径
        """
        # 加载pandoc生成的文档
        self.doc = Document(docx_path)
        
        # 保存原始markdown用于列表层级判断
        self.original_markdown = original_markdown
        
        # 设置页面格式
        self._setup_page_format()
        
        # 处理文档内容格式
        self._format_document_content(metadata)
        
        # 添加文档标题（如果有）
        if metadata.get('title'):
            self._add_document_title(metadata['title'])
        
        # 添加附件说明
        if metadata.get('attachments'):
            for attachment in metadata['attachments']:
                self._add_attachment(attachment)
        
        # 添加页码
        self._add_page_numbers()
        
        # 格式化列表
        self.format_lists()
        
        # 格式化表格
        self.format_tables()
        
        # 格式化图片
        self.format_images()
        
        # 移除图片标题段落
        self._remove_image_captions()
        
        # 保存格式化后的文档
        self.doc.save(docx_path)
        return docx_path
    
    def _setup_page_format(self):
        """设置页面格式"""
        section = self.doc.sections[0]
        
        # 设置页边距
        section.top_margin = self.config.PAGE_MARGINS['top']
        section.bottom_margin = self.config.PAGE_MARGINS['bottom']
        section.left_margin = self.config.PAGE_MARGINS['left']
        section.right_margin = self.config.PAGE_MARGINS['right']
        
        # 设置页眉页脚距离
        section.header_distance = Mm(25)
        section.footer_distance = Mm(25)
    
    def _format_document_content(self, metadata: Dict[str, Any]):
        """格式化文档内容，应用公文格式要求"""
        for paragraph in self.doc.paragraphs:
            # 跳过空段落
            if not paragraph.text.strip():
                continue
            
            # 判断段落类型并应用相应格式
            if self._is_heading(paragraph):
                self._format_heading(paragraph)
            else:
                self._format_body_paragraph(paragraph)
    
    def _is_heading(self, paragraph) -> bool:
        """判断段落是否为标题"""
        # 检查是否使用了标题样式
        if paragraph.style.name.startswith('Heading'):
            return True
        
        # 检查段落内容是否以中文数字或序号开头（如"一、"、"（一）"等）
        text = paragraph.text.strip()
        heading_patterns = [
            r'^[一二三四五六七八九十]+、',  # 一、二、三、
            r'^（[一二三四五六七八九十]+）',  # （一）（二）（三）
            r'^[0-9]+\.',  # 1. 2. 3.
            r'^[0-9]+、',  # 1、2、3、
        ]
        
        for pattern in heading_patterns:
            if re.match(pattern, text):
                return True
        
        return False
    
    def _format_heading(self, paragraph):
        """格式化标题段落"""
        text = paragraph.text.strip()
        level = self._get_heading_level(paragraph, text)
        
        # 清除现有格式
        for run in paragraph.runs:
            run.clear()
        
        # 重新添加文本
        run = paragraph.add_run(text)
        
        # 根据级别应用格式
        if level == 1:
            # 一级标题：黑体，三号，不加粗
            run.font.name = self.config.FONTS['heiti']
            run.font.size = self.config.FONT_SIZES['body']
            run.bold = False
            # 设置字体颜色为黑色
            run.font.color.rgb = RGBColor(0, 0, 0)
            self._set_chinese_font(run, self.config.FONTS['heiti'])
        elif level == 2:
            # 二级标题：楷体，三号，不加粗
            run.font.name = self.config.FONTS['kaiti']
            run.font.size = self.config.FONT_SIZES['body']
            run.bold = False
            # 设置字体颜色为黑色
            run.font.color.rgb = RGBColor(0, 0, 0)
            self._set_chinese_font(run, self.config.FONTS['kaiti'])
        else:
            # 三级及以下标题：仿宋，三号，不加粗
            run.font.name = self.config.FONTS['fangsong']
            run.font.size = self.config.FONT_SIZES['body']
            run.bold = False
            # 设置字体颜色为黑色
            run.font.color.rgb = RGBColor(0, 0, 0)
            self._set_chinese_font(run, self.config.FONTS['fangsong'])
        
        # 设置段落格式
        paragraph.alignment = self.config.ALIGNMENTS['justify']
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = self.config.FIRST_LINE_INDENT
        paragraph_format.line_spacing = self.config.LINE_SPACING
        paragraph_format.space_after = Pt(6)
        paragraph_format.space_before = Pt(6)
    
    def _get_heading_level(self, paragraph, text: str) -> int:
        """获取标题级别"""
        # 检查Word内置标题样式 - pandoc生成的样式映射（修正版）
        # 因为pandoc将#作为Heading 1，我们跳过#，所以：
        # ## → Heading 2 → 应该是我们的一级标题（黑体）
        # ### → Heading 3 → 应该是我们的二级标题（楷体）  
        # #### → Heading 4 → 应该是我们的三级标题（仿宋）
        if paragraph.style.name == 'Heading 1':
            return 1  # # → 黑体 (但我们通常跳过这个)
        elif paragraph.style.name == 'Heading 2':
            return 1  # ## → 黑体
        elif paragraph.style.name == 'Heading 3':
            return 2  # ### → 楷体
        elif paragraph.style.name == 'Heading 4':
            return 3  # #### → 仿宋
        elif paragraph.style.name == 'Heading 5':
            return 3  # ##### → 仿宋 (fallback)
        
        # 根据文本内容判断级别（处理中文标题格式）
        if re.match(r'^[一二三四五六七八九十]+、', text):
            return 1  # 一、二、三、 → 黑体
        elif re.match(r'^（[一二三四五六七八九十]+）', text):
            return 2  # （一）（二）（三） → 楷体
        elif re.match(r'^[0-9]+\.', text) or re.match(r'^[0-9]+、', text):
            return 3  # 1. 2. 3. 或 1、2、3、 → 仿宋
        
        return 3  # 默认三级标题
    
    def _format_body_paragraph(self, paragraph):
        """格式化正文段落"""
        # 为所有运行应用仿宋格式
        for run in paragraph.runs:
            run.font.name = self.config.FONTS['fangsong']
            run.font.size = self.config.FONT_SIZES['body']
            self._set_chinese_font(run, self.config.FONTS['fangsong'])
        
        # 设置段落格式
        paragraph.alignment = self.config.ALIGNMENTS['justify']
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = self.config.FIRST_LINE_INDENT
        paragraph_format.line_spacing = self.config.LINE_SPACING
        paragraph_format.space_after = Pt(0)
        paragraph_format.space_before = Pt(0)
    
    def _add_document_title(self, title: str):
        """在文档开头添加标题"""
        # 在文档开头插入标题段落
        title_paragraph = self.doc.paragraphs[0].insert_paragraph_before()
        title_paragraph.alignment = self.config.ALIGNMENTS['center']
        
        run = title_paragraph.add_run(title)
        run.font.name = self.config.FONTS['xiaobiaosong']
        run.font.size = self.config.FONT_SIZES['title']
        run.bold = True
        
        # 设置中文字体
        self._set_chinese_font(run, self.config.FONTS['xiaobiaosong'])
        
        # 添加空行
        self.doc.paragraphs[1].insert_paragraph_before()
    
    def _add_attachment(self, attachment: str):
        """添加附件说明"""
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = self.config.ALIGNMENTS['justify']
        
        run = paragraph.add_run(attachment)
        run.font.name = self.config.FONTS['fangsong']
        run.font.size = self.config.FONT_SIZES['body']
        
        self._set_chinese_font(run, self.config.FONTS['fangsong'])
        
        # 设置段落格式
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = self.config.FIRST_LINE_INDENT
        paragraph_format.line_spacing = self.config.LINE_SPACING
        paragraph_format.space_after = Pt(0)
    

    
    def _add_page_numbers(self):
        """添加页码"""
        section = self.doc.sections[0]
        footer = section.footer
        
        # 清除现有内容
        footer.paragraphs[0].clear()
        
        # 添加页码
        paragraph = footer.paragraphs[0]
        paragraph.alignment = self.config.ALIGNMENTS['center']
        
        run = paragraph.add_run("- ")
        run.font.name = self.config.FONTS['fangsong']
        run.font.size = self.config.FONT_SIZES['page_num']
        self._set_chinese_font(run, self.config.FONTS['fangsong'])
        
        # 插入页码字段
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run2 = paragraph.add_run()
        run2._r.append(fldChar1)
        run2._r.append(instrText)
        run2._r.append(fldChar2)
        run2.font.name = self.config.FONTS['fangsong']
        run2.font.size = self.config.FONT_SIZES['page_num']
        self._set_chinese_font(run2, self.config.FONTS['fangsong'])
        
        run3 = paragraph.add_run(" -")
        run3.font.name = self.config.FONTS['fangsong']
        run3.font.size = self.config.FONT_SIZES['page_num']
        self._set_chinese_font(run3, self.config.FONTS['fangsong'])
    
    def _set_chinese_font(self, run, font_name: str):
        """设置中文字体"""
        rPr = run._element.rPr
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            run._element.insert(0, rPr)
        
        # 设置中文字体
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:hint'), 'eastAsia')
    
    def format_tables(self):
        """格式化表格，包含完整的自动适应功能"""
        for table in self.doc.tables:
            # 启用表格自动适应
            if self.config.TABLE_CONFIG['auto_fit']:
                table.autofit = True
                
                # 设置表格对齐方式为居中
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # 通过XML设置表格自动适应窗口
                tbl = table._tbl
                tblPr = tbl.tblPr
                
                # 设置表格宽度为100%
                if self.config.TABLE_CONFIG['auto_fit_mode'] == 'window':
                    tblW = tblPr.xpath('.//w:tblW')[0] if tblPr.xpath('.//w:tblW') else None
                    if tblW is None:
                        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
                        tblPr.append(tblW)
                    else:
                        tblW.set(qn('w:w'), str(self.config.TABLE_CONFIG['preferred_width_percent'] * 50))
                        tblW.set(qn('w:type'), 'pct')
                
                # 设置表格布局为自动
                tblLayout = tblPr.xpath('.//w:tblLayout')[0] if tblPr.xpath('.//w:tblLayout') else None
                if tblLayout is None:
                    tblLayout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="autofit"/>')
                    tblPr.append(tblLayout)
                else:
                    tblLayout.set(qn('w:type'), 'autofit')
                
                # 设置表格允许跨页断行
                if self.config.TABLE_CONFIG['allow_row_breaks']:
                    tblPrEx = tbl.xpath('.//w:tblPrEx')[0] if tbl.xpath('.//w:tblPrEx') else None
                    if tblPrEx is None:
                        tblPrEx = parse_xml(f'<w:tblPrEx {nsdecls("w")}><w:tblLayout w:type="autofit"/></w:tblPrEx>')
                        tbl.append(tblPrEx)
            
            # 格式化表格内容
            for row in table.rows:
                # 设置行高规则
                if hasattr(row, '_tr'):
                    trPr = row._tr.xpath('.//w:trPr')[0] if row._tr.xpath('.//w:trPr') else None
                    if trPr is None:
                        trPr = parse_xml(f'<w:trPr {nsdecls("w")}></w:trPr>')
                        row._tr.insert(0, trPr)
                    
                    # 设置行高规则为自动
                    if self.config.TABLE_CONFIG['row_height_rule'] == 'auto':
                        trHeight = trPr.xpath('.//w:trHeight')[0] if trPr.xpath('.//w:trHeight') else None
                        if trHeight is None:
                            trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:hRule="auto"/>')
                            trPr.append(trHeight)
                        else:
                            trHeight.set(qn('w:hRule'), 'auto')
                
                for cell in row.cells:
                    # 设置单元格垂直对齐为居中
                    tc = cell._tc
                    tcPr = tc.xpath('.//w:tcPr')[0] if tc.xpath('.//w:tcPr') else None
                    if tcPr is None:
                        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
                        tc.insert(0, tcPr)
                    
                    vAlign = tcPr.xpath('.//w:vAlign')[0] if tcPr.xpath('.//w:vAlign') else None
                    if vAlign is None:
                        vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
                        tcPr.append(vAlign)
                    else:
                        vAlign.set(qn('w:val'), 'center')
                    
                    # 格式化单元格内的段落
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = self.config.FONTS['fangsong']
                            run.font.size = self.config.FONT_SIZES['table']  # 使用4号字体
                            self._set_chinese_font(run, self.config.FONTS['fangsong'])
                        
                        # 设置单元格段落格式
                        paragraph.alignment = self.config.ALIGNMENTS['center']  # 表格内容居中
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.line_spacing = self.config.LINE_SPACING
                        paragraph_format.space_before = Pt(3)
                        paragraph_format.space_after = Pt(3)
    
    def format_lists(self):
        """格式化列表 - 保留Word列表格式，只调整缩进和字体"""
        for i, paragraph in enumerate(self.doc.paragraphs):
            
            # 检查是否为列表项 - 检查Compact样式或有列表编号的段落
            has_list_numbering = False
            if hasattr(paragraph, '_element'):
                numPr = paragraph._element.xpath('.//w:numPr')
                has_list_numbering = bool(numPr)
            
            if paragraph.style.name == 'Compact' or has_list_numbering:
                paragraph_format = paragraph.paragraph_format
                
                # 通过列表级别判断是否为二级列表
                is_sub_item = False
                if hasattr(paragraph, '_element'):
                    ilvl = paragraph._element.xpath('.//w:ilvl/@w:val')
                    if ilvl and int(ilvl[0]) > 0:
                        is_sub_item = True
                
                if is_sub_item:
                    # 二级列表：使用配置文件中的缩进设置
                    paragraph_format.left_indent = self.config.LIST_INDENT['level2_left']
                    paragraph_format.first_line_indent = self.config.LIST_INDENT['level2_first_line']
                    paragraph_format.space_after = Pt(0)
                    paragraph_format.space_before = Pt(0)
                    paragraph.alignment = self.config.ALIGNMENTS['justify']
                else:
                    # 一级列表：使用配置文件中的悬挂缩进设置
                    paragraph_format.left_indent = self.config.LIST_INDENT['level1_left']
                    paragraph_format.first_line_indent = self.config.LIST_INDENT['level1_first_line']
                    paragraph_format.space_after = Pt(0)
                    paragraph_format.space_before = Pt(0)
                    paragraph.alignment = self.config.ALIGNMENTS['justify']
                
                # 设置字体格式（保留原有文本，只改字体）
                for run in paragraph.runs:
                    run.font.name = self.config.FONTS['fangsong']
                    run.font.size = self.config.FONT_SIZES['body']
                    self._set_chinese_font(run, self.config.FONTS['fangsong'])
            
            # 检查有序列表 - 包含数字序号的段落
            elif re.match(r'^\d+\.', paragraph.text.strip()):
                paragraph_format = paragraph.paragraph_format
                
                # 检查是否需要分割段落（包含多个列表项）
                text = paragraph.text.strip()
                items = re.split(r'(\d+\.)', text)[1:]  # 分割并保留序号
                
                if len(items) > 2:  # 包含多个列表项
                    
                    # 保存当前段落位置
                    current_p = paragraph._element
                    parent = current_p.getparent()
                    
                    # 重新组织项目
                    paired_items = []
                    for i in range(0, len(items), 2):
                        if i + 1 < len(items):
                            paired_items.append(items[i] + items[i + 1])
                    
                    # 更新第一个段落为第一项
                    if paired_items:
                        paragraph.clear()
                        first_run = paragraph.add_run(paired_items[0].strip())
                        first_run.font.name = self.config.FONTS['fangsong']
                        first_run.font.size = self.config.FONT_SIZES['body']
                        self._set_chinese_font(first_run, self.config.FONTS['fangsong'])
                        
                        paragraph_format.left_indent = Pt(0)
                        paragraph_format.first_line_indent = self.config.FIRST_LINE_INDENT
                        paragraph_format.space_after = Pt(0)
                        paragraph_format.space_before = Pt(0)
                        paragraph.alignment = self.config.ALIGNMENTS['justify']
                        
                        # 为剩余项目创建新段落
                        for item_text in paired_items[1:]:
                            new_p = paragraph._element.makeelement(paragraph._element.tag, paragraph._element.attrib)
                            parent.insert(parent.index(current_p) + 1, new_p)
                            current_p = new_p
                            
                            # 创建新的段落对象
                            from docx.text.paragraph import Paragraph
                            new_paragraph = Paragraph(new_p, paragraph._parent)
                            
                            # 添加文本
                            new_run = new_paragraph.add_run(item_text.strip())
                            new_run.font.name = self.config.FONTS['fangsong']
                            new_run.font.size = self.config.FONT_SIZES['body']
                            self._set_chinese_font(new_run, self.config.FONTS['fangsong'])
                            
                            # 设置格式
                            new_paragraph_format = new_paragraph.paragraph_format
                            new_paragraph_format.left_indent = Pt(0)
                            new_paragraph_format.first_line_indent = self.config.FIRST_LINE_INDENT
                            new_paragraph_format.space_after = Pt(0)
                            new_paragraph_format.space_before = Pt(0)
                            new_paragraph.alignment = self.config.ALIGNMENTS['justify']
                else:
                    # 单个有序列表项，直接格式化
                    paragraph_format.left_indent = Pt(0)
                    paragraph_format.first_line_indent = self.config.FIRST_LINE_INDENT
                    paragraph_format.space_after = Pt(0)
                    paragraph_format.space_before = Pt(0)
                    paragraph.alignment = self.config.ALIGNMENTS['justify']
                    
                    # 确保列表项使用正确字体
                    for run in paragraph.runs:
                        run.font.name = self.config.FONTS['fangsong']
                        run.font.size = self.config.FONT_SIZES['body']
                        self._set_chinese_font(run, self.config.FONTS['fangsong'])
    
    def format_images(self):
        """格式化文档中的图片"""
        try:
            # 遍历文档中的所有段落，查找包含图片的段落
            for paragraph in self.doc.paragraphs:
                # 检查段落中是否包含inline shapes（图片）
                if hasattr(paragraph, '_element'):
                    # 查找段落中的图片元素
                    drawings = paragraph._element.xpath('.//w:drawing')
                    for drawing in drawings:
                        self._format_single_image(drawing)
        except Exception as e:
            print(f"图片格式化时出现错误: {e}")
    
    def _format_single_image(self, drawing_element):
        """格式化单个图片"""
        try:
            # 移除图片文件名
            self._remove_image_name(drawing_element)
            
            # 设置文字环绕（如果启用）
            if self.config.PANDOC_CONFIG.get('image_wrap_text', False):
                self._set_image_wrap(drawing_element)
                
        except Exception as e:
            print(f"格式化单个图片时出现错误: {e}")
    
    def _remove_image_name(self, drawing_element):
        """移除图片的文件名显示"""
        try:
            # 查找所有可能包含图片名称的元素
            
            # 1. wp:docPr元素（Word图片文档属性）
            docPr_elements = drawing_element.xpath('.//wp:docPr')
            for docPr in docPr_elements:
                if docPr.get('title'):
                    docPr.set('title', '')  # 清空标题
                if docPr.get('descr'):
                    docPr.set('descr', '')  # 清空描述
                if docPr.get('name'):
                    docPr.set('name', '')  # 完全清空名称
            
            # 2. pic:cNvPr元素（图片核心非可视属性）
            cNvPr_elements = drawing_element.xpath('.//pic:cNvPr')
            for cNvPr in cNvPr_elements:
                if cNvPr.get('name'):
                    cNvPr.set('name', '')  # 完全清空名称
                if cNvPr.get('descr'):
                    cNvPr.set('descr', '')  # 清空描述
                if cNvPr.get('title'):
                    cNvPr.set('title', '')  # 清空标题
            
            # 3. a:blip元素（图片链接）
            blip_elements = drawing_element.xpath('.//a:blip')
            for blip in blip_elements:
                if blip.get('title'):
                    blip.set('title', '')
                if blip.get('cstate'):
                    blip.set('cstate', '')
            
            # 4. 查找并移除图片标题段落（在图片后面的文字）
            parent_paragraph = drawing_element.getparent()
            while parent_paragraph is not None and parent_paragraph.tag != qn('w:p'):
                parent_paragraph = parent_paragraph.getparent()
            
            if parent_paragraph is not None:
                # 检查图片后是否有文本内容是文件名
                for run in parent_paragraph.xpath('.//w:r'):
                    text_elements = run.xpath('.//w:t')
                    for text_elem in text_elements:
                        if text_elem.text and ('Pasted image' in text_elem.text or 
                                             '006Fd7o3gy1' in text_elem.text or
                                             '.png' in text_elem.text or
                                             '.jpg' in text_elem.text):
                            text_elem.text = ''  # 清空图片文件名文本
                    
        except Exception as e:
            print(f"移除图片名称时出现错误: {e}")
    
    def _set_image_wrap(self, drawing_element):
        """设置图片文字环绕为top and bottom"""
        try:
            if not self.config.PANDOC_CONFIG.get('image_wrap_text', False):
                return
                
            wrap_type = self.config.PANDOC_CONFIG.get('image_wrap_type', 'topAndBottom')
            
            # 查找inline元素（内联图片）- 使用更简单的xpath方法
            inlines = drawing_element.xpath('.//wp:inline')
            
            for inline in inlines:
                # 获取父元素
                parent = inline.getparent()
                
                # 获取现有的图片尺寸
                extent = inline.xpath('.//wp:extent')
                cx = extent[0].get('cx') if extent else '3000000'
                cy = extent[0].get('cy') if extent else '2000000'
                
                # 获取docPr信息
                docPr = inline.xpath('.//wp:docPr')
                docPr_id = docPr[0].get('id') if docPr else '1'
                docPr_name = docPr[0].get('name') if docPr else 'Picture'
                
                # 获取graphic元素
                graphic = inline.xpath('.//a:graphic')
                
                if graphic:
                    # 创建anchor元素替换inline元素
                    anchor_xml = f'''<wp:anchor {nsdecls("wp")} {nsdecls("a")}
                        distT="0" distB="0" distL="114300" distR="114300" 
                        simplePos="0" relativeHeight="251658240" 
                        behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">
                        <wp:simplePos x="0" y="0"/>
                        <wp:positionH relativeFrom="column">
                            <wp:align>center</wp:align>
                        </wp:positionH>
                        <wp:positionV relativeFrom="paragraph">
                            <wp:posOffset>0</wp:posOffset>
                        </wp:positionV>
                        <wp:extent cx="{cx}" cy="{cy}"/>
                        <wp:effectExtent l="0" t="0" r="0" b="0"/>
                        <wp:wrapTopAndBottom/>
                        <wp:docPr id="{docPr_id}" name="{docPr_name}"/>
                        <wp:cNvGraphicFramePr>
                            <a:graphicFrameLocks noChangeAspect="1"/>
                        </wp:cNvGraphicFramePr>
                    </wp:anchor>'''
                    
                    anchor = parse_xml(anchor_xml)
                    
                    # 复制graphic元素到新的anchor中
                    anchor.append(graphic[0])
                    
                    # 替换inline元素
                    parent.replace(inline, anchor)
                    
        except Exception as e:
            print(f"设置图片环绕时出现错误: {e}")
            # 如果出错，保持原有的inline格式
    
    def _remove_image_captions(self):
        """移除图片标题段落"""
        try:
            paragraphs_to_remove = []
            
            for i, paragraph in enumerate(self.doc.paragraphs):
                text = paragraph.text.strip()
                
                # 检查是否为图片标题段落
                if text and self._is_image_caption(text):
                    paragraphs_to_remove.append(paragraph)
                    continue
                
                # 检查是否为图片标题段落（如"图 3：MOCVD工艺流程"）
                if re.match(r'^图\s*\d+\s*[:：]\s*', text):
                    # 保留完整的图片标题，设置为caption格式
                    self._format_image_caption(paragraph)
            
            # 移除标识为需要删除的段落
            for paragraph in paragraphs_to_remove:
                self._remove_paragraph(paragraph)
                
        except Exception as e:
            print(f"移除图片标题时出现错误: {e}")
    
    def _is_image_caption(self, text: str) -> bool:
        """判断是否为需要移除的图片文件名文本"""
        # 只移除明显的文件名，不移除有意义的图片标题
        file_name_patterns = [
            r'Pasted image \d+',
            r'^006Fd7o3gy1.*\.(png|jpg|jpeg|gif|bmp)$',  # 纯文件名
            r'^Screenshot.*\.(png|jpg|jpeg|gif|bmp)$',   # 纯截图文件名
            r'^.*\.(png|jpg|jpeg|gif|bmp)$',  # 纯文件名（不包含中文描述）
        ]
        
        # 排除有意义的图片标题（包含"图 X："格式的）
        if re.match(r'^图\s*\d+\s*[:：]', text):
            return False
        
        for pattern in file_name_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                # 进一步检查是否包含中文描述
                if re.search(r'[\u4e00-\u9fff]', text) and len(text) > 20:
                    # 包含中文且较长，可能是有意义的描述
                    return False
                return True
        return False
    
    def _format_image_caption(self, paragraph):
        """格式化图片标题为caption格式"""
        try:
            # 设置图片标题的格式
            for run in paragraph.runs:
                run.font.name = self.config.FONTS['fangsong']
                run.font.size = self.config.FONT_SIZES['table']  # 使用4号字体
                run.bold = False
                self._set_chinese_font(run, self.config.FONTS['fangsong'])
            
            # 设置段落格式 - 图片标题居中显示
            paragraph.alignment = self.config.ALIGNMENTS['center']
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing = self.config.LINE_SPACING
            paragraph_format.space_after = Pt(6)
            paragraph_format.space_before = Pt(3)
            paragraph_format.first_line_indent = Pt(0)  # 图片标题不缩进
            
            # 可选：设置为Word的Caption样式（如果需要）
            # paragraph.style = 'Caption'  # 这需要确保Caption样式存在
            
        except Exception as e:
            print(f"格式化图片标题时出现错误: {e}")
    
    def _remove_paragraph(self, paragraph):
        """安全地移除段落"""
        try:
            p = paragraph._element
            p.getparent().remove(p)
            paragraph._element = None
        except Exception as e:
            print(f"移除段落时出现错误: {e}")