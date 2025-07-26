"""
图片格式化器 - 负责图片处理和格式化
"""
from docx import Document
from docx.shared import Pt
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml.parser import parse_xml
from typing import Dict, Any
import re

from ..config.config import DocumentConfig
from ..utils.xpath_cache import OptimizedXMLProcessor
from ..utils.constants import Patterns, DocumentFormats
from ..utils.exceptions import (
    ImageProcessingError,
    XMLProcessingError
)
from .base_formatter import BaseFormatter


class ImageFormatter(BaseFormatter):
    """图片格式化器 - 负责图片处理和格式化"""
    
    def __init__(self, config: DocumentConfig = None):
        super().__init__(config)
        self.xml_processor = OptimizedXMLProcessor()
    
    def format_images(self, doc: Document):
        """格式化文档中的图片 - 使用优化的单次遍历，并移除所有图片相关文件名"""
        try:
            # 1. 使用优化的方法一次性查找所有包含图片的段落
            drawings_map = self.xml_processor.find_drawings_in_paragraphs(doc.paragraphs)
            
            # 处理每个包含图片的段落
            for paragraph_index, drawings in drawings_map.items():
                for drawing in drawings:
                    self._format_single_image(drawing)
            
            # 2. 扫描所有段落，查找并移除图片文件名文本（即使没有绘制元素）
            self._remove_image_captions_from_all_paragraphs(doc)
            
            # 3. 格式化图片标题段落（处理"Image Caption"样式和"图 X："格式）
            self._format_all_image_captions(doc)
            
            # 4. 移除因删除图片标题而产生的空白段落
            self._remove_empty_paragraphs_after_image_cleanup(doc)
                    
        except (AttributeError, KeyError) as e:
            raise ImageProcessingError(f"图片元素访问错误: {e}")
        except Exception as e:
            raise ImageProcessingError(f"图片格式化时出现错误: {e}")
    
    def _format_single_image(self, drawing_element):
        """格式化单个图片"""
        try:
            # 移除图片文件名
            self._remove_image_name(drawing_element)
            
            # 设置图片宽度为全宽
            self._set_image_full_width(drawing_element)
            
            # 设置文字环绕（如果启用）
            if self.config.PANDOC_CONFIG.get('image_wrap_text', False):
                self._set_image_wrap(drawing_element)
                
        except Exception as e:
            # 记录错误但不中断处理
            import logging
            logging.debug(f"格式化单个图片时出错: {e}")
    
    def _remove_image_name(self, drawing_element):
        """移除图片的文件名显示 - 使用优化的批量处理"""
        try:
            # 使用优化的批量查询获取所有图片相关元素
            elements = self.xml_processor.process_image_properties(drawing_element)
            
            # 1. 处理wp:docPr元素（Word图片文档属性）
            for docPr in elements.get('docPr', []):
                if docPr.get('title'):
                    docPr.set('title', '')  # 清空标题
                if docPr.get('descr'):
                    docPr.set('descr', '')  # 清空描述
                if docPr.get('name'):
                    docPr.set('name', '')  # 完全清空名称
            
            # 2. 处理pic:cNvPr元素（图片核心非可视属性）
            for cNvPr in elements.get('cNvPr', []):
                if cNvPr.get('name'):
                    cNvPr.set('name', '')  # 完全清空名称
                if cNvPr.get('descr'):
                    cNvPr.set('descr', '')  # 清空描述
                if cNvPr.get('title'):
                    cNvPr.set('title', '')  # 清空标题
            
            # 3. 处理a:blip元素（图片链接）
            for blip in elements.get('blip', []):
                if blip.get('title'):
                    blip.set('title', '')
                if blip.get('cstate'):
                    blip.set('cstate', '')
            
            # 4. 查找并移除图片标题段落（在图片后面的文字）
            parent_paragraph = drawing_element.getparent()
            while parent_paragraph is not None and parent_paragraph.tag != qn('w:p'):
                parent_paragraph = parent_paragraph.getparent()
            
            if parent_paragraph is not None:
                # 检查图片后是否有文本内容是文件名
                runs = parent_paragraph.xpath('.//w:r')
                for run in runs:
                    text_elements = run.xpath('.//w:t')
                    for text_elem in text_elements:
                        if text_elem.text and ('Pasted image' in text_elem.text or 
                                             '006Fd7o3gy1' in text_elem.text or
                                             '.png' in text_elem.text or
                                             '.jpg' in text_elem.text):
                            text_elem.text = ''  # 清空图片文件名文本
                    
        except (AttributeError, KeyError) as e:
            # XML结构访问错误，记录但不中断
            import logging
            logging.debug(f"移除图片名称时XML结构访问错误: {e}")
        except Exception as e:
            import logging
            logging.debug(f"移除图片名称时出现未知错误: {e}")
    
    def _remove_image_captions_from_all_paragraphs(self, doc: Document):
        """扫描所有段落，移除图片文件名文本（不依赖绘制元素）"""
        try:
            image_filename_patterns = DocumentFormats.IMAGE_CLEANUP_PATTERNS
            
            for paragraph in doc.paragraphs:
                if not paragraph.text.strip():
                    continue
                
                # 检查段落是否包含数学公式，如果包含则跳过
                if self._has_math_formula(paragraph):
                    continue
                
                # 检查段落文本是否包含图片文件名模式
                text_contains_image_pattern = any(
                    pattern in paragraph.text for pattern in image_filename_patterns
                )
                
                if text_contains_image_pattern:
                    # 逐个检查段落中的run
                    for run in paragraph.runs:
                        if run.text:
                            original_text = run.text
                            modified_text = original_text
                            
                            # 移除包含图片文件名模式的部分，但保护数学公式
                            for pattern in image_filename_patterns:
                                if pattern in modified_text:
                                    # 如果是Pasted image模式，移除整个"Pasted image 日期时间"格式
                                    if pattern == 'Pasted image':
                                        modified_text = Patterns.PASTED_IMAGE_CLEANUP_PATTERN.sub('', modified_text)
                                    else:
                                        # 对于文件扩展名，移除包含该扩展名的词，但保护数学公式
                                        words = modified_text.split()
                                        filtered_words = []
                                        for word in words:
                                            # 检查词是否包含数学公式，如果包含则保留
                                            if (Patterns.LATEX_INLINE_MATH_PATTERN.search(word) or 
                                                Patterns.LATEX_BLOCK_MATH_PATTERN.search(word)):
                                                filtered_words.append(word)
                                            elif pattern not in word:
                                                filtered_words.append(word)
                                            # 如果词既包含文件名模式又包含数学公式，优先保护数学公式
                                        modified_text = ' '.join(filtered_words)
                            
                            # 只在文本确实改变时更新
                            if modified_text != original_text:
                                run.text = modified_text.strip()
                                
        except Exception as e:
            # 记录错误但不中断处理
            import logging
            logging.debug(f"移除图片标题文本时出错: {e}")
    
    def _format_all_image_captions(self, doc: Document):
        """格式化所有图片标题段落（包括Pandoc生成的"Image Caption"样式）"""
        try:
            for paragraph in doc.paragraphs:
                # 跳过包含数学公式的段落
                if self._has_math_formula(paragraph):
                    continue
                
                # 检查是否为Pandoc生成的图片标题样式
                if paragraph.style.name == 'Image Caption':
                    self._format_image_caption(paragraph)
                    continue
                
                # 检查是否为"图 X："格式的标题
                text = paragraph.text.strip()
                if text and Patterns.CAPTION_PATTERN.match(text):
                    self._format_image_caption(paragraph)
                    
        except Exception as e:
            # 记录错误但不中断处理
            import logging
            logging.debug(f"格式化图片标题时出错: {e}")
    
    def _remove_empty_paragraphs_after_image_cleanup(self, doc: Document):
        """移除因删除图片标题而产生的空白段落"""
        try:
            # 收集需要删除的段落（逆序遍历以避免索引变化问题）
            paragraphs_to_remove = []
            
            for i in range(len(doc.paragraphs) - 1, -1, -1):
                paragraph = doc.paragraphs[i]
                
                # 检查段落是否为空或只包含空白字符
                if self._is_effectively_empty_paragraph(paragraph):
                    paragraphs_to_remove.append(paragraph)
            
            # 删除空段落
            for paragraph in paragraphs_to_remove:
                # 获取段落的父元素并删除
                paragraph_element = paragraph._element
                parent = paragraph_element.getparent()
                if parent is not None:
                    parent.remove(paragraph_element)
                    
        except Exception as e:
            # 记录错误但不中断处理
            import logging
            logging.debug(f"移除空白段落时出错: {e}")
    
    def _set_image_full_width(self, drawing_element):
        """设置图片宽度为全宽（适用于inline元素）"""
        try:
            # 使用优化的批量查询获取所有图片相关元素
            elements = self.xml_processor.process_image_properties(drawing_element)
            
            # 计算页面可用宽度 (A4纸宽210mm - 左边距28mm - 右边距26mm = 156mm)
            # 转换为EMU单位: 1mm = 36000 EMU
            page_width_mm = 210 - 28 - 26  # 156mm
            page_width_emu = page_width_mm * 36000  # 5616000 EMU
            
            # 获取并更新extent元素
            for extent in elements.get('extent', []):
                # 获取原始尺寸
                cx_original = extent.get('cx', '3000000')
                cy_original = extent.get('cy', '2000000')
                
                # 计算纵横比并根据新宽度计算高度
                try:
                    cx_orig_int = int(cx_original)
                    cy_orig_int = int(cy_original)
                    aspect_ratio = cy_orig_int / cx_orig_int
                    
                    # 设置图片宽度为页面宽度
                    extent.set('cx', str(page_width_emu))
                    # 按比例计算新高度
                    extent.set('cy', str(int(page_width_emu * aspect_ratio)))
                except (ValueError, ZeroDivisionError):
                    # 如果转换失败，只设置宽度
                    extent.set('cx', str(page_width_emu))
            
            # 同时更新pic:spPr中的extent（如果存在）
            pic_extent_xpath = './/pic:spPr/a:xfrm/a:ext'
            pic_extents = drawing_element.xpath(pic_extent_xpath, 
                                               namespaces={'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
                                                          'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
            for pic_extent in pic_extents:
                cx_original = pic_extent.get('cx', '3000000')
                cy_original = pic_extent.get('cy', '2000000')
                
                try:
                    cx_orig_int = int(cx_original)
                    cy_orig_int = int(cy_original)
                    aspect_ratio = cy_orig_int / cx_orig_int
                    
                    pic_extent.set('cx', str(page_width_emu))
                    pic_extent.set('cy', str(int(page_width_emu * aspect_ratio)))
                except (ValueError, ZeroDivisionError):
                    pic_extent.set('cx', str(page_width_emu))
                    
        except Exception as e:
            # 记录错误但不中断处理
            import logging
            logging.debug(f"设置图片全宽时出错: {e}")
    
    def _is_effectively_empty_paragraph(self, paragraph) -> bool:
        """检查段落是否实际为空（没有文本内容或只有空白）"""
        try:
            # 检查段落文本
            if paragraph.text.strip():
                return False
            
            # 检查是否包含数学公式 - 非常重要！
            if self._has_math_formula(paragraph):
                return False
            
            # 检查是否包含非文本内容（如图片、表格等）
            if hasattr(paragraph, '_element'):
                # 检查是否有绘图元素
                drawings = paragraph._element.xpath('.//w:drawing')
                if drawings:
                    return False
                    
                # 检查是否有其他非空内容元素
                content_elements = paragraph._element.xpath('.//w:r[w:t[normalize-space(.)]]')
                if content_elements:
                    return False
                
                # 检查是否有嵌入的对象或特殊元素
                objects = paragraph._element.xpath('.//w:object | .//w:pict')
                if objects:
                    return False
            
            return True
            
        except Exception:
            # 如果检查过程出错，保守处理：不删除
            return False
    
    def _set_image_wrap(self, drawing_element):
        """设置图片文字环绕为top and bottom，并设置图片宽度为全宽"""
        try:
            if not self.config.PANDOC_CONFIG.get('image_wrap_text', False):
                return
                
            # 使用优化的批量查询获取所有图片相关元素
            elements = self.xml_processor.process_image_properties(drawing_element)
            
            for inline in elements.get('inline', []):
                # 获取父元素
                parent = inline.getparent()
                
                # 从批量查询结果中获取元素
                extent = elements.get('extent', [])
                cx_original = extent[0].get('cx') if extent else '3000000'
                cy_original = extent[0].get('cy') if extent else '2000000'
                
                # 计算页面可用宽度 (A4纸宽210mm - 左边距28mm - 右边距26mm = 156mm)
                # 转换为EMU单位: 1mm = 36000 EMU
                page_width_mm = 210 - 28 - 26  # 156mm
                page_width_emu = page_width_mm * 36000  # 5616000 EMU
                
                # 计算纵横比并根据新宽度计算高度
                try:
                    cx_orig_int = int(cx_original)
                    cy_orig_int = int(cy_original)
                    aspect_ratio = cy_orig_int / cx_orig_int
                    
                    # 设置图片宽度为页面宽度
                    cx = str(page_width_emu)
                    # 按比例计算新高度
                    cy = str(int(page_width_emu * aspect_ratio))
                except (ValueError, ZeroDivisionError):
                    # 如果转换失败，使用默认值
                    cx = str(page_width_emu)
                    cy = cy_original
                
                # 获取docPr信息
                docPr = elements.get('docPr', [])
                docPr_id = docPr[0].get('id') if docPr else '1'
                docPr_name = docPr[0].get('name') if docPr else 'Picture'
                
                # 获取graphic元素
                graphic = elements.get('graphic', [])
                
                if graphic:
                    # 安全创建anchor元素，避免XML注入漏洞
                    anchor = self._create_anchor_element(cx, cy, docPr_id, docPr_name)
                    
                    # 复制graphic元素到新的anchor中
                    anchor.append(graphic[0])
                    
                    # 替换inline元素
                    parent.replace(inline, anchor)
                    
        except (AttributeError, KeyError, Exception) as e:
            # 记录错误但不中断处理
            import logging
            logging.debug(f"设置图片文字环绕时出错: {e}")
    
    def _create_anchor_element(self, cx: str, cy: str, docPr_id: str, docPr_name: str):
        """
        安全创建anchor元素，避免XML注入漏洞
        
        Args:
            cx: 图片宽度
            cy: 图片高度  
            docPr_id: 文档属性ID
            docPr_name: 文档属性名称
            
        Returns:
            创建的anchor XML元素
        """
        try:
            # 输入验证和清理
            cx = str(cx).strip() if cx else '3000000'
            cy = str(cy).strip() if cy else '2000000'
            docPr_id = str(docPr_id).strip() if docPr_id else '1'
            docPr_name = str(docPr_name).strip() if docPr_name else 'Picture'
            
            # 验证数值参数
            try:
                int(cx)
                int(cy) 
                int(docPr_id)
            except ValueError:
                # 如果参数无效，使用默认值
                cx, cy, docPr_id = '3000000', '2000000', '1'
            
            # 使用原生XML API安全构建anchor元素
            anchor = OxmlElement('wp:anchor')
            
            # 设置anchor属性 (修复: 使用unqualified属性名，符合WordprocessingML规范)
            anchor.set('distT', '0')
            anchor.set('distB', '0') 
            anchor.set('distL', '114300')
            anchor.set('distR', '114300')
            anchor.set('simplePos', '0')
            anchor.set('relativeHeight', '251658240')
            anchor.set('behindDoc', '0')
            anchor.set('locked', '0')
            anchor.set('layoutInCell', '1')
            anchor.set('allowOverlap', '1')
            
            # 创建simplePos子元素
            simplePos = OxmlElement('wp:simplePos')
            simplePos.set('x', '0')
            simplePos.set('y', '0')
            anchor.append(simplePos)
            
            # 创建positionH子元素
            positionH = OxmlElement('wp:positionH')
            positionH.set('relativeFrom', 'column')
            align = OxmlElement('wp:align')
            align.text = 'center'
            positionH.append(align)
            anchor.append(positionH)
            
            # 创建positionV子元素
            positionV = OxmlElement('wp:positionV')
            positionV.set('relativeFrom', 'paragraph')
            posOffset = OxmlElement('wp:posOffset')
            posOffset.text = '0'
            positionV.append(posOffset)
            anchor.append(positionV)
            
            # 创建extent子元素
            extent = OxmlElement('wp:extent')
            extent.set('cx', cx)
            extent.set('cy', cy)
            anchor.append(extent)
            
            # 创建effectExtent子元素
            effectExtent = OxmlElement('wp:effectExtent')
            effectExtent.set('l', '0')
            effectExtent.set('t', '0')
            effectExtent.set('r', '0')
            effectExtent.set('b', '0')
            anchor.append(effectExtent)
            
            # 创建wrapTopAndBottom子元素
            wrapTopAndBottom = OxmlElement('wp:wrapTopAndBottom')
            anchor.append(wrapTopAndBottom)
            
            # 创建docPr子元素
            docPr = OxmlElement('wp:docPr')
            docPr.set('id', docPr_id)
            docPr.set('name', docPr_name)
            anchor.append(docPr)
            
            # 创建cNvGraphicFramePr子元素
            cNvGraphicFramePr = OxmlElement('wp:cNvGraphicFramePr')
            graphicFrameLocks = OxmlElement('a:graphicFrameLocks')
            graphicFrameLocks.set('noChangeAspect', '1')
            cNvGraphicFramePr.append(graphicFrameLocks)
            anchor.append(cNvGraphicFramePr)
            
            return anchor
            
        except (XMLProcessingError, Exception) as e:
            # 记录错误并返回基本anchor元素
            import logging
            logging.debug(f"创建图片anchor元素时出错: {e}")
            fallback_anchor = OxmlElement('wp:anchor')
            return fallback_anchor
    
    def _is_image_caption(self, text: str) -> bool:
        """判断是否为需要移除的图片文件名文本"""
        # 排除有意义的图片标题（包含"图 X："格式的）
        if Patterns.CAPTION_PATTERN.match(text):
            return False
        
        for pattern in Patterns.IMAGE_FILENAME_PATTERNS:
            if pattern.search(text, re.IGNORECASE):
                # 进一步检查是否包含中文描述
                if Patterns.CHINESE_CHAR_PATTERN.search(text) and len(text) > 20:
                    # 包含中文且较长，可能是有意义的描述
                    return False
                return True
        return False
    
    def _format_image_caption(self, paragraph):
        """格式化图片标题为caption格式"""
        try:
            # 设置图片标题的格式
            for run in paragraph.runs:
                run.font.name = self.config.FONTS['fangsong']
                run.font.size = self.config.FONT_SIZES['table']  # 使用4号字体
                run.bold = False
                self._set_chinese_font(run, self.config.FONTS['fangsong'])
            
            # 设置段落格式 - 图片标题居中显示
            paragraph.alignment = self.config.ALIGNMENTS['center']
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(0)
            paragraph_format.space_before = Pt(0)
            paragraph_format.first_line_indent = Pt(0)  # 图片标题不缩进
            
        except (AttributeError, Exception) as e:
            # 记录错误但不中断处理
            import logging
            logging.debug(f"格式化图片标题段落时出错: {e}")
    
    def _remove_paragraph(self, paragraph):
        """安全地移除段落"""
        try:
            p = paragraph._element
            p.getparent().remove(p)
            paragraph._element = None
        except (AttributeError, ValueError, Exception) as e:
            # 记录错误但不中断处理
            import logging
            logging.debug(f"移除段落时出错: {e}")