"""
文档标题格式化器 - 负责文档标题和附件处理
"""
from docx import Document
from docx.shared import Pt
from .base_formatter import BaseFormatter


class DocumentTitleFormatter(BaseFormatter):
    """文档标题格式化器 - 负责文档标题和附件处理"""
    
    def add_document_title(self, doc: Document, title: str):
        """在文档开头添加标题"""
        # 在文档开头插入标题段落
        title_paragraph = doc.paragraphs[0].insert_paragraph_before()
        title_paragraph.alignment = self.config.ALIGNMENTS['center']
        
        run = title_paragraph.add_run(title)
        run.font.name = self.config.FONTS['xiaobiaosong']
        run.font.size = self.config.FONT_SIZES['title']
        run.bold = True
        
        # 设置中文字体
        self._set_chinese_font(run, self.config.FONTS['xiaobiaosong'])
        
        # 添加空行
        doc.paragraphs[1].insert_paragraph_before()
    
    def add_attachment(self, doc: Document, attachment: str):
        """添加附件说明"""
        paragraph = doc.add_paragraph()
        paragraph.alignment = self.config.ALIGNMENTS['justify']
        
        run = paragraph.add_run(attachment)
        run.font.name = self.config.FONTS['fangsong']
        run.font.size = self.config.FONT_SIZES['body']
        
        self._set_chinese_font(run, self.config.FONTS['fangsong'])
        
        # 设置段落格式
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = self.config.FIRST_LINE_INDENT
        paragraph_format.line_spacing = self.config.LINE_SPACING
        paragraph_format.space_after = Pt(0)