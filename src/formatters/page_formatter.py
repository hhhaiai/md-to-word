"""
页面格式化器 - 负责页面设置、页眉页脚、页码等
"""
from docx import Document
from docx.shared import Mm, Pt
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import qn as qn_func
from .base_formatter import BaseFormatter


class PageFormatter(BaseFormatter):
    """页面格式化器 - 负责页面设置、页眉页脚、页码等"""
    
    def setup_page_format(self, doc: Document):
        """设置页面格式"""
        section = doc.sections[0]
        
        # 设置页边距
        section.top_margin = self.config.PAGE_MARGINS['top']
        section.bottom_margin = self.config.PAGE_MARGINS['bottom']
        section.left_margin = self.config.PAGE_MARGINS['left']
        section.right_margin = self.config.PAGE_MARGINS['right']
        
        # 设置页眉页脚距离
        section.header_distance = Mm(25)
        section.footer_distance = Mm(25)
        
        # 设置文档网格（版芯）
        self._setup_document_grid(section)
    
    def add_page_numbers(self, doc: Document):
        """添加页码"""
        section = doc.sections[0]
        footer = section.footer
        
        # 清除现有内容
        footer.paragraphs[0].clear()
        
        # 添加页码
        paragraph = footer.paragraphs[0]
        paragraph.alignment = self.config.ALIGNMENTS['center']
        
        run = paragraph.add_run("- ")
        run.font.name = self.config.FONTS['fangsong']
        run.font.size = self.config.FONT_SIZES['page_num']
        self._set_chinese_font(run, self.config.FONTS['fangsong'])
        
        # 插入页码字段
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run2 = paragraph.add_run()
        run2._r.append(fldChar1)
        run2._r.append(instrText)
        run2._r.append(fldChar2)
        run2.font.name = self.config.FONTS['fangsong']
        run2.font.size = self.config.FONT_SIZES['page_num']
        self._set_chinese_font(run2, self.config.FONTS['fangsong'])
        
        run3 = paragraph.add_run(" -")
        run3.font.name = self.config.FONTS['fangsong']
        run3.font.size = self.config.FONT_SIZES['page_num']
        self._set_chinese_font(run3, self.config.FONTS['fangsong'])
    
    def _setup_document_grid(self, section):
        """设置文档网格以强制每页22行、每行28字"""
        sectPr = section._sectPr
        
        # 检查是否已有docGrid元素
        docGrid = sectPr.find(qn_func('w:docGrid'))
        if docGrid is None:
            # 创建新的docGrid元素
            docGrid = OxmlElement('w:docGrid')
            sectPr.append(docGrid)
        
        # 设置文档网格类型为"行和字符网格"
        docGrid.set(qn_func('w:type'), 'linesAndChars')
        
        # 设置行距和字符间距
        # 行距：26.5磅 = 530 twips（1磅 = 20 twips）
        line_pitch = 530  # twips
        
        # 字符间距：版芯宽度156mm ÷ 28字 ≈ 5.57mm ≈ 316 twips
        char_space = 316  # twips
        
        # 设置行距（单位：twips，1/20点）
        docGrid.set(qn_func('w:linePitch'), str(line_pitch))
        
        # 设置字符间距（单位：twips）
        docGrid.set(qn_func('w:charSpace'), str(char_space))