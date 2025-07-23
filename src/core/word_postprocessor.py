from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any, List
import os
import re

from ..config.config import DocumentConfig
from ..utils.constants import Patterns
from ..formatters import (
    PageFormatter, 
    ParagraphFormatter, 
    DocumentTitleFormatter, 
    TableFormatter, 
    ListFormatter, 
    ImageFormatter
)


class WordPostprocessor:
    """
    重构后的Word文档后处理器
    使用组合模式，将不同的格式化功能委托给专门的格式化器类
    解决了原有的"God Object"反模式问题
    
    特别处理：包含数学公式的图片caption
    - 检测包含LaTeX数学公式的caption（如 $\\theta$, $\\omega$）
    - 使用特殊的分离处理逻辑，避免破坏MathML内容
    - 通过元素属性标记已处理的caption，防止重复处理
    """
    
    def __init__(self):
        self.config = DocumentConfig()
        self.doc = None
        
        # 初始化专门的格式化器
        self.page_formatter = PageFormatter(self.config)
        self.paragraph_formatter = ParagraphFormatter(self.config)
        self.title_formatter = DocumentTitleFormatter(self.config)
        self.table_formatter = TableFormatter(self.config)
        self.list_formatter = ListFormatter(self.config)
        self.image_formatter = ImageFormatter(self.config)
    
    def apply_formatting(self, docx_path: str, metadata: Dict[str, Any], original_markdown: str = None) -> str:
        """
        对pandoc生成的Word文档应用公文格式
        
        Args:
            docx_path: pandoc生成的Word文档路径
            metadata: 包含标题、附件等元数据的字典
            original_markdown: 原始markdown内容，用于判断列表层级
            
        Returns:
            处理后的Word文档路径
        """
        # 加载pandoc生成的文档
        self.doc = Document(docx_path)
        
        # 保存原始markdown用于列表层级判断
        self.original_markdown = original_markdown
        
        # 保存文档路径的目录，用于查找图片
        self.doc_dir = os.path.dirname(os.path.abspath(docx_path))
        
        # 使用专门的格式化器处理不同方面的格式化
        self.page_formatter.setup_page_format(self.doc)
        self.paragraph_formatter.format_document_content(self.doc, metadata)
        
        # 添加文档标题（如果有）
        if metadata.get('title'):
            self.title_formatter.add_document_title(self.doc, metadata['title'])
        
        # 添加附件说明
        if metadata.get('attachments'):
            for attachment in metadata['attachments']:
                self.title_formatter.add_attachment(self.doc, attachment)
        
        # 应用各种格式化
        self.page_formatter.add_page_numbers(self.doc)
        self.list_formatter.format_lists(self.doc)
        self.table_formatter.format_tables(self.doc)
        
        # 新的图片处理方式：直接查找并替换图片语法
        self.process_and_insert_images()
        
        # 保存格式化后的文档
        self.doc.save(docx_path)
        return docx_path
    
    def _has_math_formula(self, paragraph) -> bool:
        """检查段落是否包含数学公式"""
        try:
            # 1. 检查段落的XML内容是否包含MathML元素（转换后的数学公式）
            xml_str = paragraph._element.xml
            if 'oMath' in xml_str or 'oMathPara' in xml_str:
                return True
            
            # 2. 检查段落文本是否包含LaTeX格式的数学公式（原始格式）
            text = paragraph.text
            if text:
                # 检查行内数学公式 $...$
                if Patterns.LATEX_INLINE_MATH_PATTERN.search(text):
                    return True
                # 检查块级数学公式 $$...$$
                if Patterns.LATEX_BLOCK_MATH_PATTERN.search(text):
                    return True
            
            return False
        except:
            return False
    
    # 保留原有的公共方法以维持向后兼容性
    def format_tables(self):
        """格式化表格（向后兼容方法）"""
        if self.doc:
            self.table_formatter.format_tables(self.doc)
    
    def format_lists(self):
        """格式化列表（向后兼容方法）"""
        if self.doc:
            self.list_formatter.format_lists(self.doc)
    
    def format_images(self):
        """格式化图片（向后兼容方法）"""
        if self.doc:
            self.image_formatter.format_images(self.doc)
    
    def process_and_insert_images(self):
        """处理文档中的图片语法并插入实际图片"""
        # 初始化已处理的数学caption文本集合
        self._processed_math_caption_texts = set()
        
        # 定义图片语法模式
        markdown_image_pattern = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')
        obsidian_image_pattern = re.compile(r'!\[\[([^\]]+)\]\]')
        
        # 综合的图表标题模式：匹配 图/图片/表/表格/图表 + 可选空格 + 数字 + 可选空格 + 标点(:：.) + 描述
        caption_pattern = re.compile(r'^(图片?|表格?|图表)\s*(\d+)\s*[:：.]\s*')
        
        # 图片计数器
        image_counter = 0
        # 收集需要处理的图片信息
        images_to_process = []
        
        # 第一遍：识别所有图片和caption
        all_paragraphs = list(self.doc.paragraphs)
        for i, paragraph in enumerate(all_paragraphs):
            text = paragraph.text.strip()
            
            # 检查是否包含图片语法
            markdown_match = markdown_image_pattern.search(text)
            obsidian_match = obsidian_image_pattern.search(text)
            
            if markdown_match or obsidian_match:
                image_counter += 1
                
                # 检查是否在同一段落包含caption
                caption_in_same_para = None
                if markdown_match or obsidian_match:
                    # 提取图片语法后的文本作为potential caption
                    if markdown_match:
                        image_syntax_end = markdown_match.end()
                        remaining_text = text[image_syntax_end:].strip()
                        alt_text = markdown_match.group(1)
                        image_path = markdown_match.group(2)
                        image_type = 'markdown'
                    else:
                        image_syntax_end = obsidian_match.end()
                        remaining_text = text[image_syntax_end:].strip()
                        image_path = obsidian_match.group(1)
                        # 确定标题
                        if 'Pasted image' in image_path:
                            alt_text = ""  # 空标题，不显示
                        else:
                            alt_text = os.path.splitext(os.path.basename(image_path))[0]
                        image_type = 'obsidian'
                    
                    # 检查剩余文本是否为caption
                    if remaining_text and caption_pattern.match(remaining_text):
                        caption_in_same_para = remaining_text
                
                # 检查下一段落是否为caption（跳过空行）
                caption_paragraph = None
                for j in range(i + 1, min(i + 3, len(all_paragraphs))):  # 最多检查后面2个段落
                    next_text = all_paragraphs[j].text.strip()
                    if not next_text:  # 跳过空段落
                        continue
                    if caption_pattern.match(next_text):
                        caption_paragraph = all_paragraphs[j]
                        break
                    else:
                        break  # 遇到非空、非caption段落就停止
                
                # 查找图片实际路径
                actual_path = self._find_image_actual_path(image_path)
                
                # 构建图片信息
                image_info = {
                    'path': actual_path,
                    'title': alt_text,
                    'type': image_type,
                    'original': text,
                    'number': image_counter,
                    'paragraph': paragraph,
                    'caption_in_same_para': caption_in_same_para,
                    'caption_paragraph': caption_paragraph
                }
                
                images_to_process.append(image_info)
        
        # 第二遍：处理图片插入
        for image_info in images_to_process:
            self._replace_paragraph_with_image(image_info['paragraph'], image_info)
        
        # 第三遍：处理caption格式化
        self._process_captions()
    
    def _find_image_actual_path(self, image_path: str) -> str:
        """查找图片的实际路径"""
        # 如果是绝对路径或URL，直接返回
        if os.path.isabs(image_path) or image_path.startswith(('http://', 'https://')):
            return image_path
        
        # 在配置的搜索路径中查找图片
        from pathlib import Path
        
        # 构建搜索路径列表
        search_paths = [
            self.doc_dir,  # 首先在源文件目录查找
            *DocumentConfig.get_image_search_paths()  # 然后在配置的搜索路径中查找
        ]
        
        # 支持的图片格式
        supported_formats = DocumentConfig.IMAGE_CONFIG['supported_formats']
        
        for search_path_str in search_paths:
            search_path = Path(search_path_str).resolve()
            if not search_path.exists() or not search_path.is_dir():
                continue
            
            try:
                # 构建目标图片路径
                image_path_obj = (search_path / image_path).resolve()
                
                # 直接匹配文件名
                if image_path_obj.is_file():
                    return str(image_path_obj)
                
                # 如果没有扩展名，尝试添加支持的格式
                if not image_path_obj.suffix:
                    for ext in supported_formats:
                        path_with_ext = image_path_obj.with_suffix(ext)
                        if path_with_ext.is_file():
                            return str(path_with_ext)
                            
            except (OSError, ValueError):
                continue
        
        return None
    
    def _replace_paragraph_with_image(self, paragraph, image_info: Dict):
        """将包含图片语法的段落替换为实际图片"""
        try:
            # 检查图片路径是否存在
            if not image_info['path'] or not os.path.exists(image_info['path']):
                # 如果图片不存在，保留原文本
                return
            
            # 检查段落是否包含数学公式
            has_math = self._has_math_formula(paragraph)
            
            # 如果包含数学公式，需要特殊处理（优先处理）
            if has_math:
                # 分离处理：在当前段落前插入图片段落，保留原段落作为caption
                self._insert_image_before_math_caption(paragraph, image_info)
                return
            
            # 处理同一段落中的caption分离（只在没有数学公式时处理）
            if image_info.get('caption_in_same_para'):
                # 在当前段落后插入caption段落
                p_element = paragraph._element
                parent = p_element.getparent()
                
                # 创建caption段落
                caption_p = self.doc.add_paragraph()
                caption_p.text = image_info['caption_in_same_para']
                
                # 将caption段落移动到图片段落之后
                parent.insert(parent.index(p_element) + 1, caption_p._element)
            
            # 如果没有数学公式，可以安全地清空段落
            paragraph.clear()
            
            # 添加图片到段落
            run = paragraph.add_run()
            picture = run.add_picture(image_info['path'], width=Inches(5))  # 默认宽度5英寸
            
            # 设置段落居中
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 设置图片文字环绕
            drawing_elements = run._element.xpath('.//w:drawing')
            if drawing_elements:
                self.image_formatter._set_image_wrap(drawing_elements[0])
                
        except Exception as e:
            # 保留原始文本
            paragraph.text = image_info['original']
    
    def _remove_image_syntax_only(self, paragraph, image_info: Dict):
        """只移除图片语法部分，保留其他内容（包括数学公式）"""
        try:
            import re
            
            # 获取原始文本
            original_text = paragraph.text
            
            # 移除图片语法部分
            # 移除 Obsidian 格式 ![[filename]]
            cleaned_text = re.sub(r'!\[\[[^\]]+\]\]', '', original_text)
            # 移除 Markdown 格式 ![alt](path)
            cleaned_text = re.sub(r'!\[[^\]]*\]\([^)]+\)', '', cleaned_text)
            
            # 清理多余的空格
            cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()
            
            # 只有当清理后的文本确实不同时才更新
            if cleaned_text != original_text:
                
                # 需要更精细地处理，避免破坏MathML
                # 遍历段落中的运行，只修改文本部分
                for run in paragraph.runs:
                    if run.text:
                        # 检查这个运行是否包含图片语法
                        if '![[' in run.text or '![' in run.text:
                            # 清理这个运行中的图片语法
                            new_run_text = re.sub(r'!\[\[[^\]]+\]\]', '', run.text)
                            new_run_text = re.sub(r'!\[[^\]]*\]\([^)]+\)', '', new_run_text)
                            new_run_text = re.sub(r'\s+', ' ', new_run_text).strip()
                            
                            if new_run_text != run.text:
                                run.text = new_run_text
                                
        except Exception as e:
            pass  # 静默处理错误
    
    def _insert_image_before_math_caption(self, caption_paragraph, image_info: Dict):
        """在包含数学公式的caption前插入图片段落"""
        try:
            # 创建新的图片段落
            p_element = caption_paragraph._element
            parent = p_element.getparent()
            
            # 在当前段落前创建图片段落
            image_p = self.doc.add_paragraph()
            
            # 添加图片到新段落
            run = image_p.add_run()
            picture = run.add_picture(image_info['path'], width=Inches(5))
            
            # 设置图片段落居中
            image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 设置图片文字环绕
            drawing_elements = run._element.xpath('.//w:drawing')
            if drawing_elements:
                self.image_formatter._set_image_wrap(drawing_elements[0])
            
            # 将图片段落移动到caption段落前面
            parent.insert(parent.index(p_element), image_p._element)
            
            # 清理caption段落中的图片语法，但保留数学公式
            self._remove_image_syntax_only(caption_paragraph, image_info)
            
            # 记录处理后的caption文本，用于跳过重复处理
            cleaned_text = caption_paragraph.text.strip()
            if cleaned_text:
                self._processed_math_caption_texts.add(cleaned_text)
            
            # 直接格式化这个caption，因为它包含数学公式
            self.image_formatter._format_image_caption(caption_paragraph)
            
        except Exception as e:
            # 降级处理：只移除图片语法
            self._remove_image_syntax_only(caption_paragraph, image_info)
    
    def _process_captions(self):
        """格式化所有图片和表格caption（位置已在预处理阶段调整）"""
        try:
            # caption识别模式
            caption_pattern = re.compile(r'^(图片?|图表|表格?)\s*(\d+)\s*[.:：]\s*(.*?)$')
            
            # 格式化所有caption，不再处理位置
            for paragraph in self.doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # 检查是否是已处理的数学caption，跳过避免重复处理
                if hasattr(self, '_processed_math_caption_texts') and text in self._processed_math_caption_texts:
                    continue
                
                # 检查是否匹配caption模式
                caption_match = caption_pattern.match(text)
                if caption_match:
                    caption_type = caption_match.group(1)
                    number = caption_match.group(2)
                    content = caption_match.group(3)
                    
                    # 检查是否包含数学公式
                    has_math = self._has_math_formula(paragraph)
                    
                    # 标准化格式（仅在没有数学公式时替换文本）
                    if not has_math:
                        if caption_type in ['图', '图片', '图表']:
                            paragraph.text = f"图{number}. {content}"
                        else:
                            paragraph.text = f"表{number}. {content}"
                    
                    # 应用格式
                    if caption_type in ['图', '图片', '图表']:
                        self.image_formatter._format_image_caption(paragraph)
                    else:
                        self._format_table_caption(paragraph)
            
        except Exception as e:
            pass  # 静默处理错误
    
    def _format_table_caption(self, paragraph):
        """格式化表格caption，与图片caption相同的格式"""
        try:
            # 使用与图片caption相同的格式设置
            for run in paragraph.runs:
                run.font.name = self.config.FONTS['fangsong']
                run.font.size = self.config.FONT_SIZES['table']  # 4号字体
                run.bold = False
                # 设置中文字体
                from docx.oxml.ns import qn
                run._element.rPr.rFonts.set(qn('w:eastAsia'), self.config.FONTS['fangsong'])
            
            # 设置段落格式 - 表格caption居中显示
            paragraph.alignment = self.config.ALIGNMENTS['center']
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing = self.config.LINE_SPACING
            paragraph_format.space_after = Pt(6)
            paragraph_format.space_before = Pt(3)
            paragraph_format.first_line_indent = Pt(0)  # caption不缩进
            
        except AttributeError as e:
            pass  # 静默处理属性错误
        except Exception as e:
            pass  # 静默处理其他错误
    
