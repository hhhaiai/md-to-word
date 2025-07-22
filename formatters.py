"""
专业格式化器类 - 将WordPostprocessor的功能分解为单一职责的格式化器
"""
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.oxml.shared import OxmlElement, qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml.parser import parse_xml
from typing import Dict, Any
import re

from config import DocumentConfig
from xpath_cache import OptimizedXMLProcessor
from constants import Patterns, DocumentFormats
from exceptions import (
    DocumentFormattingError, 
    ImageProcessingError, 
    TableFormattingError, 
    ListFormattingError,
    XMLProcessingError
)


class BaseFormatter:
    """格式化器基类，提供共同的功能和配置"""
    
    def __init__(self, config: DocumentConfig = None):
        self.config = config or DocumentConfig()
    
    def _set_chinese_font(self, run, font_name: str):
        """设置中文字体"""
        rPr = run._element.rPr
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            run._element.insert(0, rPr)
        
        # 设置中文字体
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:hint'), 'eastAsia')
    
    def _has_math_formula(self, paragraph) -> bool:
        """检查段落是否包含数学公式（统一方法，避免重复）"""
        try:
            # 1. 检查段落的XML内容是否包含MathML元素（转换后的数学公式）
            xml_str = paragraph._element.xml
            if 'oMath' in xml_str or 'oMathPara' in xml_str:
                return True
            
            # 2. 检查段落文本是否包含LaTeX格式的数学公式（原始格式）
            text = paragraph.text
            if text:
                # 检查行内数学公式 $...$
                if Patterns.LATEX_INLINE_MATH_PATTERN.search(text):
                    return True
                # 检查块级数学公式 $$...$$
                if Patterns.LATEX_BLOCK_MATH_PATTERN.search(text):
                    return True
            
            return False
        except:
            return False


class PageFormatter(BaseFormatter):
    """页面格式化器 - 负责页面设置、页眉页脚、页码等"""
    
    def setup_page_format(self, doc: Document):
        """设置页面格式"""
        section = doc.sections[0]
        
        # 设置页边距
        section.top_margin = self.config.PAGE_MARGINS['top']
        section.bottom_margin = self.config.PAGE_MARGINS['bottom']
        section.left_margin = self.config.PAGE_MARGINS['left']
        section.right_margin = self.config.PAGE_MARGINS['right']
        
        # 设置页眉页脚距离
        section.header_distance = Mm(25)
        section.footer_distance = Mm(25)
    
    def add_page_numbers(self, doc: Document):
        """添加页码"""
        section = doc.sections[0]
        footer = section.footer
        
        # 清除现有内容
        footer.paragraphs[0].clear()
        
        # 添加页码
        paragraph = footer.paragraphs[0]
        paragraph.alignment = self.config.ALIGNMENTS['center']
        
        run = paragraph.add_run("- ")
        run.font.name = self.config.FONTS['fangsong']
        run.font.size = self.config.FONT_SIZES['page_num']
        self._set_chinese_font(run, self.config.FONTS['fangsong'])
        
        # 插入页码字段
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run2 = paragraph.add_run()
        run2._r.append(fldChar1)
        run2._r.append(instrText)
        run2._r.append(fldChar2)
        run2.font.name = self.config.FONTS['fangsong']
        run2.font.size = self.config.FONT_SIZES['page_num']
        self._set_chinese_font(run2, self.config.FONTS['fangsong'])
        
        run3 = paragraph.add_run(" -")
        run3.font.name = self.config.FONTS['fangsong']
        run3.font.size = self.config.FONT_SIZES['page_num']
        self._set_chinese_font(run3, self.config.FONTS['fangsong'])


class ParagraphFormatter(BaseFormatter):
    """段落格式化器 - 负责标题和正文段落的格式化"""
    
    def format_document_content(self, doc: Document, metadata: Dict[str, Any]):
        """格式化文档内容，应用公文格式要求"""
        for paragraph in doc.paragraphs:
            # 跳过空段落
            if not paragraph.text.strip():
                continue
            
            # 判断段落类型并应用相应格式
            if self._is_heading(paragraph):
                self._format_heading(paragraph)
            else:
                self._format_body_paragraph(paragraph)
    
    def _is_heading(self, paragraph) -> bool:
        """判断段落是否为标题"""
        # 检查是否使用了标题样式
        if paragraph.style.name.startswith('Heading'):
            return True
        
        # 检查段落内容是否以中文数字或序号开头（如"一、"、"（一）"等）
        text = paragraph.text.strip()
        
        for pattern in Patterns.HEADING_PATTERNS:
            if pattern.match(text):
                return True
        
        return False
    
    def _format_heading(self, paragraph):
        """格式化标题段落"""
        text = paragraph.text.strip()
        level = self._get_heading_level(paragraph, text)
        
        # 检查是否包含数学公式（通过检查XML元素）
        if self._has_math_formula(paragraph):
            # 如果包含数学公式，只修改字体格式，不清除内容
            self._format_paragraph_with_math(paragraph, level, is_heading=True)
            return
        
        # 对于不包含数学公式的标题，使用原有的方法
        # 清除现有格式
        for run in paragraph.runs:
            run.clear()
        
        # 重新添加文本
        run = paragraph.add_run(text)
        
        # 根据级别应用格式
        if level == 1:
            # 一级标题：黑体，三号，不加粗
            run.font.name = self.config.FONTS['heiti']
            run.font.size = self.config.FONT_SIZES['body']
            run.bold = False
            # 设置字体颜色为黑色
            run.font.color.rgb = RGBColor(0, 0, 0)
            self._set_chinese_font(run, self.config.FONTS['heiti'])
        elif level == 2:
            # 二级标题：楷体，三号，不加粗
            run.font.name = self.config.FONTS['kaiti']
            run.font.size = self.config.FONT_SIZES['body']
            run.bold = False
            # 设置字体颜色为黑色
            run.font.color.rgb = RGBColor(0, 0, 0)
            self._set_chinese_font(run, self.config.FONTS['kaiti'])
        else:
            # 三级及以下标题：仿宋，三号，不加粗
            run.font.name = self.config.FONTS['fangsong']
            run.font.size = self.config.FONT_SIZES['body']
            run.bold = False
            # 设置字体颜色为黑色
            run.font.color.rgb = RGBColor(0, 0, 0)
            self._set_chinese_font(run, self.config.FONTS['fangsong'])
        
        # 设置段落格式
        paragraph.alignment = self.config.ALIGNMENTS['justify']
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = self.config.FIRST_LINE_INDENT
        paragraph_format.line_spacing = self.config.LINE_SPACING
        paragraph_format.space_after = Pt(6)
        paragraph_format.space_before = Pt(6)
    
    def _get_heading_level(self, paragraph, text: str) -> int:
        """获取标题级别"""
        # 检查Word内置标题样式 - pandoc生成的样式映射（修正版）
        # 因为pandoc将#作为Heading 1，我们跳过#，所以：
        # ## → Heading 2 → 应该是我们的一级标题（黑体）
        # ### → Heading 3 → 应该是我们的二级标题（楷体）  
        # #### → Heading 4 → 应该是我们的三级标题（仿宋）
        if paragraph.style.name == 'Heading 1':
            return 1  # # → 黑体 (但我们通常跳过这个)
        elif paragraph.style.name == 'Heading 2':
            return 1  # ## → 黑体
        elif paragraph.style.name == 'Heading 3':
            return 2  # ### → 楷体
        elif paragraph.style.name == 'Heading 4':
            return 3  # #### → 仿宋
        elif paragraph.style.name == 'Heading 5':
            return 3  # ##### → 仿宋 (fallback)
        
        # 根据文本内容判断级别（处理中文标题格式）
        if Patterns.HEADING_PATTERNS[0].match(text):
            return 1  # 一、二、三、 → 黑体
        elif Patterns.HEADING_PATTERNS[1].match(text):
            return 2  # （一）（二）（三） → 楷体
        elif Patterns.HEADING_PATTERNS[2].match(text) or Patterns.HEADING_PATTERNS[3].match(text):
            return 3  # 1. 2. 3. 或 1、2、3、 → 仿宋
        
        return 3  # 默认三级标题
    
    def _format_body_paragraph(self, paragraph):
        """格式化正文段落"""
        # 检查是否包含数学公式
        if self._has_math_formula(paragraph):
            # 如果包含数学公式，使用特殊的格式化方法
            self._format_paragraph_with_math(paragraph, level=0, is_heading=False)
            return
        
        # 为所有运行应用仿宋格式
        for run in paragraph.runs:
            run.font.name = self.config.FONTS['fangsong']
            run.font.size = self.config.FONT_SIZES['body']
            self._set_chinese_font(run, self.config.FONTS['fangsong'])
        
        # 设置段落格式
        paragraph.alignment = self.config.ALIGNMENTS['justify']
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = self.config.FIRST_LINE_INDENT
        paragraph_format.line_spacing = self.config.LINE_SPACING
        paragraph_format.space_after = Pt(0)
        paragraph_format.space_before = Pt(0)
    
    def _format_paragraph_with_math(self, paragraph, level: int, is_heading: bool):
        """格式化包含数学公式的段落，保留数学公式内容"""
        # 设置段落格式
        paragraph.alignment = self.config.ALIGNMENTS['justify']
        paragraph_format = paragraph.paragraph_format
        
        if is_heading:
            paragraph_format.first_line_indent = self.config.FIRST_LINE_INDENT
            paragraph_format.space_after = Pt(6)
            paragraph_format.space_before = Pt(6)
        else:
            paragraph_format.first_line_indent = self.config.FIRST_LINE_INDENT
            paragraph_format.space_after = Pt(0)
            paragraph_format.space_before = Pt(0)
        
        paragraph_format.line_spacing = self.config.LINE_SPACING
        
        # 只格式化文本run，跳过数学公式
        for run in paragraph.runs:
            if run._element.tag.endswith('r'):  # 普通文本run
                try:
                    # 检查run的XML内容，只处理不包含数学内容的run
                    if 'oMath' not in run._element.xml:
                        if is_heading:
                            # 根据标题级别设置字体
                            if level == 1:
                                run.font.name = self.config.FONTS['heiti']
                                self._set_chinese_font(run, self.config.FONTS['heiti'])
                            elif level == 2:
                                run.font.name = self.config.FONTS['kaiti']
                                self._set_chinese_font(run, self.config.FONTS['kaiti'])
                            else:
                                run.font.name = self.config.FONTS['fangsong']
                                self._set_chinese_font(run, self.config.FONTS['fangsong'])
                            run.font.size = self.config.FONT_SIZES['body']
                            run.bold = False
                            run.font.color.rgb = RGBColor(0, 0, 0)
                        else:
                            # 正文段落格式
                            run.font.name = self.config.FONTS['fangsong']
                            run.font.size = self.config.FONT_SIZES['body']
                            self._set_chinese_font(run, self.config.FONTS['fangsong'])
                except:
                    # 如果出错，跳过这个run
                    continue


class DocumentTitleFormatter(BaseFormatter):
    """文档标题格式化器 - 负责文档标题和附件处理"""
    
    def add_document_title(self, doc: Document, title: str):
        """在文档开头添加标题"""
        # 在文档开头插入标题段落
        title_paragraph = doc.paragraphs[0].insert_paragraph_before()
        title_paragraph.alignment = self.config.ALIGNMENTS['center']
        
        run = title_paragraph.add_run(title)
        run.font.name = self.config.FONTS['xiaobiaosong']
        run.font.size = self.config.FONT_SIZES['title']
        run.bold = True
        
        # 设置中文字体
        self._set_chinese_font(run, self.config.FONTS['xiaobiaosong'])
        
        # 添加空行
        doc.paragraphs[1].insert_paragraph_before()
    
    def add_attachment(self, doc: Document, attachment: str):
        """添加附件说明"""
        paragraph = doc.add_paragraph()
        paragraph.alignment = self.config.ALIGNMENTS['justify']
        
        run = paragraph.add_run(attachment)
        run.font.name = self.config.FONTS['fangsong']
        run.font.size = self.config.FONT_SIZES['body']
        
        self._set_chinese_font(run, self.config.FONTS['fangsong'])
        
        # 设置段落格式
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = self.config.FIRST_LINE_INDENT
        paragraph_format.line_spacing = self.config.LINE_SPACING
        paragraph_format.space_after = Pt(0)


class TableFormatter(BaseFormatter):
    """表格格式化器 - 负责表格格式化和样式设置"""
    
    def __init__(self, config: DocumentConfig = None):
        super().__init__(config)
        self.xml_processor = OptimizedXMLProcessor()
    
    def format_tables(self, doc: Document):
        """格式化表格，包含完整的自动适应功能"""
        for table in doc.tables:
            # 启用表格自动适应
            if self.config.TABLE_CONFIG['auto_fit']:
                table.autofit = True
                
                # 设置表格对齐方式为居中
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # 通过XML设置表格自动适应窗口
                tbl = table._tbl
                tblPr = tbl.tblPr
                
                # 使用优化的批量查询获取所有表格属性
                elements = self.xml_processor.process_table_properties(tbl)
                
                # 设置表格宽度为100%
                if self.config.TABLE_CONFIG['auto_fit_mode'] == 'window':
                    tblW = elements.get('tblW')
                    if tblW is None:
                        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
                        tblPr.append(tblW)
                    else:
                        tblW.set(qn('w:w'), str(self.config.TABLE_CONFIG['preferred_width_percent'] * 50))
                        tblW.set(qn('w:type'), 'pct')
                
                # 设置表格布局为自动
                tblLayout = elements.get('tblLayout')
                if tblLayout is None:
                    tblLayout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="autofit"/>')
                    tblPr.append(tblLayout)
                else:
                    tblLayout.set(qn('w:type'), 'autofit')
                
                # 设置表格允许跨页断行
                if self.config.TABLE_CONFIG['allow_row_breaks']:
                    tblPrEx = elements.get('tblPrEx')
                    if tblPrEx is None:
                        tblPrEx = parse_xml(f'<w:tblPrEx {nsdecls("w")}><w:tblLayout w:type="autofit"/></w:tblPrEx>')
                        tbl.append(tblPrEx)
            
            # 应用三线表样式
            self._apply_three_line_table_style(table)
            
            # 格式化表格内容
            for row_index, row in enumerate(table.rows):
                # 使用优化的行属性处理
                row_props = self.xml_processor.process_row_properties(row)
                if row_props:
                    trPr = row_props['trPr']
                    if trPr is None:
                        trPr = parse_xml(f'<w:trPr {nsdecls("w")}></w:trPr>')
                        row._tr.insert(0, trPr)
                    
                    # 设置行高规则为自动
                    if self.config.TABLE_CONFIG['row_height_rule'] == 'auto':
                        trHeight = row_props['trHeight']
                        if trHeight is None:
                            trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:hRule="auto"/>')
                            trPr.append(trHeight)
                        else:
                            trHeight.set(qn('w:hRule'), 'auto')
                
                for cell in row.cells:
                    # 使用优化的单元格属性处理
                    cell_props = self.xml_processor.process_cell_properties(cell)
                    tcPr = cell_props['tcPr']
                    vAlign = cell_props['vAlign']
                    
                    if vAlign is None:
                        vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
                        tcPr.append(vAlign)
                    else:
                        vAlign.set(qn('w:val'), 'center')
                    
                    # 应用三线表单元格边框
                    self._apply_three_line_cell_borders(cell, row_index, len(table.rows))
                    
                    # 格式化单元格内的段落
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = self.config.FONTS['fangsong']
                            run.font.size = self.config.FONT_SIZES['table']  # 使用4号字体
                            self._set_chinese_font(run, self.config.FONTS['fangsong'])
                        
                        # 设置单元格段落格式
                        paragraph.alignment = self.config.ALIGNMENTS['center']  # 表格内容居中
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.line_spacing = self.config.LINE_SPACING
                        paragraph_format.space_before = Pt(3)
                        paragraph_format.space_after = Pt(3)
    
    def _apply_three_line_table_style(self, table):
        """应用三线表样式 - 清除默认边框"""
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        # 移除表格默认边框
        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is not None:
            tblPr.remove(tblBorders)
        
        # 设置表格无边框样式
        no_border_xml = f'''<w:tblBorders {nsdecls("w")}>
            <w:top w:val="none" w:sz="0"/>
            <w:left w:val="none" w:sz="0"/>
            <w:bottom w:val="none" w:sz="0"/>
            <w:right w:val="none" w:sz="0"/>
            <w:insideH w:val="none" w:sz="0"/>
            <w:insideV w:val="none" w:sz="0"/>
        </w:tblBorders>'''
        
        new_borders = parse_xml(no_border_xml)
        tblPr.append(new_borders)
    
    def _apply_three_line_cell_borders(self, cell, row_index, total_rows):
        """为单元格应用三线表边框样式"""
        tc = cell._tc
        tcPr = tc.tcPr
        
        if tcPr is None:
            tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
            tc.insert(0, tcPr)
        
        # 移除现有边框设置
        existing_borders = tcPr.find(qn('w:tcBorders'))
        if existing_borders is not None:
            tcPr.remove(existing_borders)
        
        # 根据行位置设置边框
        if row_index == 0:
            # 第一行（表头）：顶部1.5磅黑色边框 + 底部0.75磅边框
            borders_xml = f'''<w:tcBorders {nsdecls("w")}>
                <w:top w:val="single" w:sz="21" w:color="000000"/>
                <w:bottom w:val="single" w:sz="9" w:color="000000"/>
                <w:left w:val="none" w:sz="0"/>
                <w:right w:val="none" w:sz="0"/>
            </w:tcBorders>'''
        elif row_index == total_rows - 1:
            # 最后一行：底部1.5磅黑色边框
            borders_xml = f'''<w:tcBorders {nsdecls("w")}>
                <w:top w:val="none" w:sz="0"/>
                <w:bottom w:val="single" w:sz="21" w:color="000000"/>
                <w:left w:val="none" w:sz="0"/>
                <w:right w:val="none" w:sz="0"/>
            </w:tcBorders>'''
        else:
            # 中间行：无边框
            borders_xml = f'''<w:tcBorders {nsdecls("w")}>
                <w:top w:val="none" w:sz="0"/>
                <w:bottom w:val="none" w:sz="0"/>
                <w:left w:val="none" w:sz="0"/>
                <w:right w:val="none" w:sz="0"/>
            </w:tcBorders>'''
        
        new_borders = parse_xml(borders_xml)
        tcPr.append(new_borders)


class ListFormatter(BaseFormatter):
    """列表格式化器 - 负责有序和无序列表的格式化"""
    
    def __init__(self, config: DocumentConfig = None):
        super().__init__(config)
        self.xml_processor = OptimizedXMLProcessor()
    
    def format_lists(self, doc: Document):
        """格式化列表 - 保留Word列表格式，只调整缩进和字体"""
        for paragraph in doc.paragraphs:
            list_type = self._detect_list_type(paragraph)
            
            if list_type == 'word_list':
                self._format_word_list_item(paragraph)
            elif list_type == 'ordered_list':
                self._format_ordered_list_paragraph(paragraph)
    
    def _detect_list_type(self, paragraph):
        """检测段落的列表类型"""
        # 检查是否为Word内置列表（Compact样式或有列表编号）
        if self._is_word_list_item(paragraph):
            return 'word_list'
        
        # 检查是否为有序列表（包含数字序号的段落）
        if Patterns.ORDERED_LIST_PATTERN.match(paragraph.text.strip()):
            return 'ordered_list'
            
        return 'none'
    
    def _is_word_list_item(self, paragraph) -> bool:
        """判断是否为Word内置列表项"""
        # 检查Compact样式
        if paragraph.style.name == 'Compact':
            return True
            
        # 检查是否有列表编号
        if hasattr(paragraph, '_element'):
            numPr = self.xml_processor.cache.find_first(paragraph._element, './/w:numPr')
            return numPr is not None
            
        return False
    
    def _format_word_list_item(self, paragraph):
        """格式化Word内置列表项"""
        is_sub_item = self._is_sub_list_item(paragraph)
        
        # 应用缩进和格式设置
        if is_sub_item:
            self._apply_sub_item_format(paragraph)
        else:
            self._apply_main_item_format(paragraph)
        
        # 设置字体格式
        self._apply_list_font_format(paragraph)
    
    def _is_sub_list_item(self, paragraph) -> bool:
        """判断是否为二级列表项"""
        if hasattr(paragraph, '_element'):
            ilvl_elem = self.xml_processor.cache.find_first(paragraph._element, './/w:ilvl')
            if ilvl_elem is not None:
                val = ilvl_elem.get(qn('w:val'))
                return bool(val and int(val) > 0)
        return False
    
    def _apply_sub_item_format(self, paragraph):
        """应用二级列表格式"""
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = self.config.LIST_INDENT['level2_left']
        paragraph_format.first_line_indent = self.config.LIST_INDENT['level2_first_line']
        paragraph_format.space_after = Pt(0)
        paragraph_format.space_before = Pt(0)
        paragraph.alignment = self.config.ALIGNMENTS['justify']
    
    def _apply_main_item_format(self, paragraph):
        """应用一级列表格式"""
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = self.config.LIST_INDENT['level1_left']
        paragraph_format.first_line_indent = self.config.LIST_INDENT['level1_first_line']
        paragraph_format.space_after = Pt(0)
        paragraph_format.space_before = Pt(0)
        paragraph.alignment = self.config.ALIGNMENTS['justify']
    
    def _apply_list_font_format(self, paragraph):
        """应用列表字体格式"""
        for run in paragraph.runs:
            run.font.name = self.config.FONTS['fangsong']
            run.font.size = self.config.FONT_SIZES['body']
            self._set_chinese_font(run, self.config.FONTS['fangsong'])
    
    def _format_ordered_list_paragraph(self, paragraph):
        """格式化有序列表段落"""
        paragraph_format = paragraph.paragraph_format
        
        # 检查是否需要分割段落（包含多个列表项）
        text = paragraph.text.strip()
        items = Patterns.ORDERED_LIST_SPLIT_PATTERN.split(text)[1:]  # 分割并保留序号
        
        if len(items) > 2:  # 包含多个列表项
            self._split_multi_item_paragraph(paragraph, items)
        else:
            # 单个有序列表项，直接格式化
            self._format_single_ordered_item(paragraph)
    
    def _split_multi_item_paragraph(self, paragraph, items):
        """分割包含多个列表项的段落"""
        paired_items = self._organize_list_items(items)
        
        if not paired_items:
            return
            
        # 更新第一个段落为第一项
        self._update_first_paragraph(paragraph, paired_items[0])
        
        # 为剩余项目创建新段落
        self._create_additional_paragraphs(paragraph, paired_items[1:])
    
    def _organize_list_items(self, items):
        """将分割的列表项重新组织为完整的项目"""
        paired_items = []
        for i in range(0, len(items), 2):
            if i + 1 < len(items):
                paired_items.append(items[i] + items[i + 1])
        return paired_items
    
    def _update_first_paragraph(self, paragraph, first_item_text):
        """更新第一个段落的内容和格式"""
        paragraph.clear()
        first_run = paragraph.add_run(first_item_text.strip())
        self._apply_list_run_format(first_run)
        self._set_ordered_list_format(paragraph)
    
    def _create_additional_paragraphs(self, original_paragraph, remaining_items):
        """为剩余的列表项创建新段落"""
        current_p = original_paragraph._element
        parent = current_p.getparent()
        
        for item_text in remaining_items:
            # 创建新的XML段落元素
            new_p = original_paragraph._element.makeelement(
                original_paragraph._element.tag, 
                original_paragraph._element.attrib
            )
            parent.insert(parent.index(current_p) + 1, new_p)
            current_p = new_p
            
            # 创建新的段落对象并设置内容
            new_paragraph = self._create_paragraph_object(new_p, original_paragraph._parent)
            self._set_paragraph_content_and_format(new_paragraph, item_text.strip())
    
    def _create_paragraph_object(self, xml_element, parent):
        """创建新的段落对象"""
        from docx.text.paragraph import Paragraph
        return Paragraph(xml_element, parent)
    
    def _set_paragraph_content_and_format(self, paragraph, text):
        """设置段落内容和格式"""
        new_run = paragraph.add_run(text)
        self._apply_list_run_format(new_run)
        self._set_ordered_list_format(paragraph)
    
    def _apply_list_run_format(self, run):
        """应用列表运行的字体格式"""
        run.font.name = self.config.FONTS['fangsong']
        run.font.size = self.config.FONT_SIZES['body']
        self._set_chinese_font(run, self.config.FONTS['fangsong'])
    
    def _format_single_ordered_item(self, paragraph):
        """格式化单个有序列表项"""
        self._set_ordered_list_format(paragraph)
        
        # 确保列表项使用正确字体
        for run in paragraph.runs:
            self._apply_list_run_format(run)
    
    def _set_ordered_list_format(self, paragraph):
        """设置有序列表的段落格式"""
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = Pt(0)
        paragraph_format.first_line_indent = self.config.FIRST_LINE_INDENT
        paragraph_format.space_after = Pt(0)
        paragraph_format.space_before = Pt(0)
        paragraph.alignment = self.config.ALIGNMENTS['justify']


class ImageFormatter(BaseFormatter):
    """图片格式化器 - 负责图片处理和格式化"""
    
    def __init__(self, config: DocumentConfig = None):
        super().__init__(config)
        self.xml_processor = OptimizedXMLProcessor()
    
    def format_images(self, doc: Document):
        """格式化文档中的图片 - 使用优化的单次遍历，并移除所有图片相关文件名"""
        try:
            # 1. 使用优化的方法一次性查找所有包含图片的段落
            drawings_map = self.xml_processor.find_drawings_in_paragraphs(doc.paragraphs)
            
            # 处理每个包含图片的段落
            for paragraph_index, drawings in drawings_map.items():
                for drawing in drawings:
                    self._format_single_image(drawing)
            
            # 2. 扫描所有段落，查找并移除图片文件名文本（即使没有绘制元素）
            self._remove_image_captions_from_all_paragraphs(doc)
            
            # 3. 格式化图片标题段落（处理"Image Caption"样式和"图 X："格式）
            self._format_all_image_captions(doc)
            
            # 4. 移除因删除图片标题而产生的空白段落
            self._remove_empty_paragraphs_after_image_cleanup(doc)
                    
        except (AttributeError, KeyError) as e:
            raise ImageProcessingError(f"图片元素访问错误: {e}")
        except Exception as e:
            raise ImageProcessingError(f"图片格式化时出现错误: {e}")
    
    def _format_single_image(self, drawing_element):
        """格式化单个图片"""
        try:
            # 移除图片文件名
            self._remove_image_name(drawing_element)
            
            # 设置文字环绕（如果启用）
            if self.config.PANDOC_CONFIG.get('image_wrap_text', False):
                self._set_image_wrap(drawing_element)
                
        except Exception as e:
            # 单个图片格式化失败不应中断整个处理流程
            pass  # 静默处理错误
    
    def _remove_image_name(self, drawing_element):
        """移除图片的文件名显示 - 使用优化的批量处理"""
        try:
            # 使用优化的批量查询获取所有图片相关元素
            elements = self.xml_processor.process_image_properties(drawing_element)
            
            # 1. 处理wp:docPr元素（Word图片文档属性）
            for docPr in elements.get('docPr', []):
                if docPr.get('title'):
                    docPr.set('title', '')  # 清空标题
                if docPr.get('descr'):
                    docPr.set('descr', '')  # 清空描述
                if docPr.get('name'):
                    docPr.set('name', '')  # 完全清空名称
            
            # 2. 处理pic:cNvPr元素（图片核心非可视属性）
            for cNvPr in elements.get('cNvPr', []):
                if cNvPr.get('name'):
                    cNvPr.set('name', '')  # 完全清空名称
                if cNvPr.get('descr'):
                    cNvPr.set('descr', '')  # 清空描述
                if cNvPr.get('title'):
                    cNvPr.set('title', '')  # 清空标题
            
            # 3. 处理a:blip元素（图片链接）
            for blip in elements.get('blip', []):
                if blip.get('title'):
                    blip.set('title', '')
                if blip.get('cstate'):
                    blip.set('cstate', '')
            
            # 4. 查找并移除图片标题段落（在图片后面的文字）
            parent_paragraph = drawing_element.getparent()
            while parent_paragraph is not None and parent_paragraph.tag != qn('w:p'):
                parent_paragraph = parent_paragraph.getparent()
            
            if parent_paragraph is not None:
                # 检查图片后是否有文本内容是文件名
                runs = parent_paragraph.xpath('.//w:r')
                for run in runs:
                    text_elements = run.xpath('.//w:t')
                    for text_elem in text_elements:
                        if text_elem.text and ('Pasted image' in text_elem.text or 
                                             '006Fd7o3gy1' in text_elem.text or
                                             '.png' in text_elem.text or
                                             '.jpg' in text_elem.text):
                            text_elem.text = ''  # 清空图片文件名文本
                    
        except (AttributeError, KeyError) as e:
            # XML结构访问错误，记录但不中断
            pass  # 静默处理结构错误
        except Exception as e:
            pass  # 静默处理错误
    
    def _remove_image_captions_from_all_paragraphs(self, doc: Document):
        """扫描所有段落，移除图片文件名文本（不依赖绘制元素）"""
        try:
            image_filename_patterns = DocumentFormats.IMAGE_CLEANUP_PATTERNS
            
            for paragraph in doc.paragraphs:
                if not paragraph.text.strip():
                    continue
                
                # 检查段落是否包含数学公式，如果包含则跳过
                if self._has_math_formula(paragraph):
                    continue
                
                # 检查段落文本是否包含图片文件名模式
                text_contains_image_pattern = any(
                    pattern in paragraph.text for pattern in image_filename_patterns
                )
                
                if text_contains_image_pattern:
                    # 逐个检查段落中的run
                    for run in paragraph.runs:
                        if run.text:
                            original_text = run.text
                            modified_text = original_text
                            
                            # 移除包含图片文件名模式的部分，但保护数学公式
                            for pattern in image_filename_patterns:
                                if pattern in modified_text:
                                    # 如果是Pasted image模式，移除整个"Pasted image 日期时间"格式
                                    if pattern == 'Pasted image':
                                        import re
                                        modified_text = re.sub(r'Pasted image \d{14}', '', modified_text)
                                    else:
                                        # 对于文件扩展名，移除包含该扩展名的词，但保护数学公式
                                        words = modified_text.split()
                                        filtered_words = []
                                        for word in words:
                                            # 检查词是否包含数学公式，如果包含则保留
                                            if (Patterns.LATEX_INLINE_MATH_PATTERN.search(word) or 
                                                Patterns.LATEX_BLOCK_MATH_PATTERN.search(word)):
                                                filtered_words.append(word)
                                            elif pattern not in word:
                                                filtered_words.append(word)
                                            # 如果词既包含文件名模式又包含数学公式，优先保护数学公式
                                        modified_text = ' '.join(filtered_words)
                            
                            # 只在文本确实改变时更新
                            if modified_text != original_text:
                                run.text = modified_text.strip()
                                
        except Exception as e:
            pass  # 静默处理错误
    
    def _format_all_image_captions(self, doc: Document):
        """格式化所有图片标题段落（包括Pandoc生成的"Image Caption"样式）"""
        try:
            for paragraph in doc.paragraphs:
                # 跳过包含数学公式的段落
                if self._has_math_formula(paragraph):
                    continue
                
                # 检查是否为Pandoc生成的图片标题样式
                if paragraph.style.name == 'Image Caption':
                    self._format_image_caption(paragraph)
                    continue
                
                # 检查是否为"图 X："格式的标题
                text = paragraph.text.strip()
                if text and Patterns.IMAGE_CAPTION_PATTERN.match(text):
                    self._format_image_caption(paragraph)
                    
        except Exception as e:
            pass  # 静默处理错误
    
    def _remove_empty_paragraphs_after_image_cleanup(self, doc: Document):
        """移除因删除图片标题而产生的空白段落"""
        try:
            # 收集需要删除的段落（逆序遍历以避免索引变化问题）
            paragraphs_to_remove = []
            
            for i in range(len(doc.paragraphs) - 1, -1, -1):
                paragraph = doc.paragraphs[i]
                
                # 检查段落是否为空或只包含空白字符
                if self._is_effectively_empty_paragraph(paragraph):
                    paragraphs_to_remove.append(paragraph)
            
            # 删除空段落
            for paragraph in paragraphs_to_remove:
                # 获取段落的父元素并删除
                paragraph_element = paragraph._element
                parent = paragraph_element.getparent()
                if parent is not None:
                    parent.remove(paragraph_element)
                    
        except Exception as e:
            pass  # 静默处理错误
    
    def _is_effectively_empty_paragraph(self, paragraph) -> bool:
        """检查段落是否实际为空（没有文本内容或只有空白）"""
        try:
            # 检查段落文本
            if paragraph.text.strip():
                return False
            
            # 检查是否包含数学公式 - 非常重要！
            if self._has_math_formula(paragraph):
                return False
            
            # 检查是否包含非文本内容（如图片、表格等）
            if hasattr(paragraph, '_element'):
                # 检查是否有绘图元素
                drawings = paragraph._element.xpath('.//w:drawing')
                if drawings:
                    return False
                    
                # 检查是否有其他非空内容元素
                content_elements = paragraph._element.xpath('.//w:r[w:t[normalize-space(.)]]')
                if content_elements:
                    return False
                
                # 检查是否有嵌入的对象或特殊元素
                objects = paragraph._element.xpath('.//w:object | .//w:pict')
                if objects:
                    return False
            
            return True
            
        except Exception:
            # 如果检查过程出错，保守处理：不删除
            return False
    
    def _set_image_wrap(self, drawing_element):
        """设置图片文字环绕为top and bottom"""
        try:
            if not self.config.PANDOC_CONFIG.get('image_wrap_text', False):
                return
                
            # wrap_type = self.config.PANDOC_CONFIG.get('image_wrap_type', 'topAndBottom')  # Currently only topAndBottom is supported
            
            # 使用优化的批量查询获取所有图片相关元素
            elements = self.xml_processor.process_image_properties(drawing_element)
            
            for inline in elements.get('inline', []):
                # 获取父元素
                parent = inline.getparent()
                
                # 从批量查询结果中获取元素
                extent = elements.get('extent', [])
                cx = extent[0].get('cx') if extent else '3000000'
                cy = extent[0].get('cy') if extent else '2000000'
                
                # 获取docPr信息
                docPr = elements.get('docPr', [])
                docPr_id = docPr[0].get('id') if docPr else '1'
                docPr_name = docPr[0].get('name') if docPr else 'Picture'
                
                # 获取graphic元素
                graphic = elements.get('graphic', [])
                
                if graphic:
                    # 安全创建anchor元素，避免XML注入漏洞
                    anchor = self._create_anchor_element(cx, cy, docPr_id, docPr_name)
                    
                    # 复制graphic元素到新的anchor中
                    anchor.append(graphic[0])
                    
                    # 替换inline元素
                    parent.replace(inline, anchor)
                    
        except (AttributeError, KeyError) as e:
            pass  # 静默处理结构错误
            # 如果出错，保持原有的inline格式
        except Exception as e:
            pass  # 静默处理错误
            # 如果出错，保持原有的inline格式
    
    def _create_anchor_element(self, cx: str, cy: str, docPr_id: str, docPr_name: str):
        """
        安全创建anchor元素，避免XML注入漏洞
        
        Args:
            cx: 图片宽度
            cy: 图片高度  
            docPr_id: 文档属性ID
            docPr_name: 文档属性名称
            
        Returns:
            创建的anchor XML元素
        """
        try:
            # 输入验证和清理
            cx = str(cx).strip() if cx else '3000000'
            cy = str(cy).strip() if cy else '2000000'
            docPr_id = str(docPr_id).strip() if docPr_id else '1'
            docPr_name = str(docPr_name).strip() if docPr_name else 'Picture'
            
            # 验证数值参数
            try:
                int(cx)
                int(cy) 
                int(docPr_id)
            except ValueError:
                # 如果参数无效，使用默认值
                cx, cy, docPr_id = '3000000', '2000000', '1'
            
            # 使用原生XML API安全构建anchor元素
            anchor = OxmlElement('wp:anchor')
            
            # 设置anchor属性 (修复: 使用unqualified属性名，符合WordprocessingML规范)
            anchor.set('distT', '0')
            anchor.set('distB', '0') 
            anchor.set('distL', '114300')
            anchor.set('distR', '114300')
            anchor.set('simplePos', '0')
            anchor.set('relativeHeight', '251658240')
            anchor.set('behindDoc', '0')
            anchor.set('locked', '0')
            anchor.set('layoutInCell', '1')
            anchor.set('allowOverlap', '1')
            
            # 创建simplePos子元素
            simplePos = OxmlElement('wp:simplePos')
            simplePos.set('x', '0')
            simplePos.set('y', '0')
            anchor.append(simplePos)
            
            # 创建positionH子元素
            positionH = OxmlElement('wp:positionH')
            positionH.set('relativeFrom', 'column')
            align = OxmlElement('wp:align')
            align.text = 'center'
            positionH.append(align)
            anchor.append(positionH)
            
            # 创建positionV子元素
            positionV = OxmlElement('wp:positionV')
            positionV.set('relativeFrom', 'paragraph')
            posOffset = OxmlElement('wp:posOffset')
            posOffset.text = '0'
            positionV.append(posOffset)
            anchor.append(positionV)
            
            # 创建extent子元素
            extent = OxmlElement('wp:extent')
            extent.set('cx', cx)
            extent.set('cy', cy)
            anchor.append(extent)
            
            # 创建effectExtent子元素
            effectExtent = OxmlElement('wp:effectExtent')
            effectExtent.set('l', '0')
            effectExtent.set('t', '0')
            effectExtent.set('r', '0')
            effectExtent.set('b', '0')
            anchor.append(effectExtent)
            
            # 创建wrapTopAndBottom子元素
            wrapTopAndBottom = OxmlElement('wp:wrapTopAndBottom')
            anchor.append(wrapTopAndBottom)
            
            # 创建docPr子元素
            docPr = OxmlElement('wp:docPr')
            docPr.set('id', docPr_id)
            docPr.set('name', docPr_name)
            anchor.append(docPr)
            
            # 创建cNvGraphicFramePr子元素
            cNvGraphicFramePr = OxmlElement('wp:cNvGraphicFramePr')
            graphicFrameLocks = OxmlElement('a:graphicFrameLocks')
            graphicFrameLocks.set('noChangeAspect', '1')
            cNvGraphicFramePr.append(graphicFrameLocks)
            anchor.append(cNvGraphicFramePr)
            
            return anchor
            
        except XMLProcessingError as e:
            pass  # 静默处理XML错误
            # 返回简单的anchor元素作为后备
            fallback_anchor = OxmlElement('wp:anchor')
            return fallback_anchor
        except Exception as e:
            pass  # 静默处理错误
            # 返回简单的anchor元素作为后备
            fallback_anchor = OxmlElement('wp:anchor')
            return fallback_anchor
    
    def remove_image_captions(self, doc: Document):
        """移除图片标题段落"""
        try:
            paragraphs_to_remove = []
            
            for paragraph in doc.paragraphs:
                # 检查段落是否包含数学公式，如果包含则跳过
                if self._has_math_formula(paragraph):
                    continue
                    
                text = paragraph.text.strip()
                
                # 检查是否为图片标题段落
                if text and self._is_image_caption(text):
                    paragraphs_to_remove.append(paragraph)
                    continue
                
                # 检查是否为图片标题段落（如"图 3：MOCVD工艺流程"）
                if Patterns.IMAGE_CAPTION_PATTERN.match(text):
                    # 保留完整的图片标题，设置为caption格式
                    self._format_image_caption(paragraph)
            
            # 移除标识为需要删除的段落
            for paragraph in paragraphs_to_remove:
                self._remove_paragraph(paragraph)
                
        except (AttributeError, IndexError) as e:
            pass  # 静默处理文档结构错误
        except Exception as e:
            pass  # 静默处理错误
    
    def _is_image_caption(self, text: str) -> bool:
        """判断是否为需要移除的图片文件名文本"""
        # 排除有意义的图片标题（包含"图 X："格式的）
        if Patterns.IMAGE_CAPTION_PATTERN.match(text):
            return False
        
        for pattern in Patterns.IMAGE_FILENAME_PATTERNS:
            if pattern.search(text, re.IGNORECASE):
                # 进一步检查是否包含中文描述
                if Patterns.CHINESE_CHAR_PATTERN.search(text) and len(text) > 20:
                    # 包含中文且较长，可能是有意义的描述
                    return False
                return True
        return False
    
    def _format_image_caption(self, paragraph):
        """格式化图片标题为caption格式"""
        try:
            # 设置图片标题的格式
            for run in paragraph.runs:
                run.font.name = self.config.FONTS['fangsong']
                run.font.size = self.config.FONT_SIZES['table']  # 使用4号字体
                run.bold = False
                self._set_chinese_font(run, self.config.FONTS['fangsong'])
            
            # 设置段落格式 - 图片标题居中显示
            paragraph.alignment = self.config.ALIGNMENTS['center']
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing = self.config.LINE_SPACING
            paragraph_format.space_after = Pt(6)
            paragraph_format.space_before = Pt(3)
            paragraph_format.first_line_indent = Pt(0)  # 图片标题不缩进
            
        except AttributeError as e:
            pass  # 静默处理属性错误
        except Exception as e:
            pass  # 静默处理错误
    
    def _remove_paragraph(self, paragraph):
        """安全地移除段落"""
        try:
            p = paragraph._element
            p.getparent().remove(p)
            paragraph._element = None
        except (AttributeError, ValueError) as e:
            pass  # 静默处理结构错误
        except Exception as e:
            pass  # 静默处理错误