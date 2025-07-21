from docx import Document
from docx.shared import Inches, Pt, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from typing import Dict, Any, List
import re

from config import DocumentConfig

class WordGenerator:
    """Word文档生成器，专门用于生成符合公文格式的Word文档"""
    
    def __init__(self):
        self.config = DocumentConfig()
        self.doc = None
    
    def create_document(self, parsed_data: Dict[str, Any], output_path: str):
        """创建Word文档"""
        self.doc = Document()
        
        # 设置页面格式
        self._setup_page_format()
        
        # 添加标题
        if parsed_data['title']:
            self._add_title(parsed_data['title'])
        
        # 添加主送机关
        if parsed_data['sender']:
            self._add_sender(parsed_data['sender'])
        
        # 添加正文内容
        for item in parsed_data['body']:
            if item['type'] == 'title':
                self._add_content_title(item['content'])
            elif item['type'] == 'heading':
                self._add_heading(item['content'], item['level'])
            elif item['type'] == 'paragraph':
                self._add_paragraph(item['content'])
        
        # 添加附件说明
        for attachment in parsed_data['attachments']:
            self._add_attachment(attachment)
        
        # 添加成文日期
        if parsed_data['date']:
            self._add_date(parsed_data['date'])
        
        # 添加页码
        self._add_page_numbers()
        
        # 保存文档
        self.doc.save(output_path)
    
    def _setup_page_format(self):
        """设置页面格式"""
        section = self.doc.sections[0]
        
        # 设置页边距
        section.top_margin = self.config.PAGE_MARGINS['top']
        section.bottom_margin = self.config.PAGE_MARGINS['bottom']
        section.left_margin = self.config.PAGE_MARGINS['left']
        section.right_margin = self.config.PAGE_MARGINS['right']
        
        # 设置页眉页脚距离
        section.header_distance = Mm(25)
        section.footer_distance = Mm(25)
    
    def _add_title(self, title: str):
        """添加标题"""
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = self.config.ALIGNMENTS['center']
        
        run = paragraph.add_run(title)
        run.font.name = self.config.FONTS['xiaobiaosong']
        run.font.size = self.config.FONT_SIZES['title']
        run.bold = True
        
        # 设置中文字体
        self._set_chinese_font(run, self.config.FONTS['xiaobiaosong'])
        
        # 添加空行
        self.doc.add_paragraph()
    
    def _add_sender(self, sender: str):
        """添加主送机关"""
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = self.config.ALIGNMENTS['justify']  # 两端对齐
        
        run = paragraph.add_run(sender)
        run.font.name = self.config.FONTS['fangsong']
        run.font.size = self.config.FONT_SIZES['body']
        
        self._set_chinese_font(run, self.config.FONTS['fangsong'])
        
        # 设置段落格式
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = Pt(0)  # 主送机关不缩进
        paragraph_format.line_spacing = self.config.LINE_SPACING
        paragraph_format.space_after = Pt(0)
        
        # 添加空行
        self.doc.add_paragraph()
    
    def _add_content_title(self, title: str):
        """添加正文中的标题（#对应的标题：小标宋，二号，居中）"""
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = self.config.ALIGNMENTS['center']
        
        run = paragraph.add_run(title)
        run.font.name = self.config.FONTS['xiaobiaosong']
        run.font.size = self.config.FONT_SIZES['title']
        run.bold = True
        
        # 设置中文字体
        self._set_chinese_font(run, self.config.FONTS['xiaobiaosong'])
        
        # 设置段落格式
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = self.config.LINE_SPACING
        paragraph_format.space_after = Pt(12)
        paragraph_format.space_before = Pt(12)
    
    def _add_paragraph(self, content: str):
        """添加正文段落"""
        if not content.strip():
            return
            
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = self.config.ALIGNMENTS['justify']
        
        run = paragraph.add_run(content)
        run.font.name = self.config.FONTS['fangsong']
        run.font.size = self.config.FONT_SIZES['body']
        
        self._set_chinese_font(run, self.config.FONTS['fangsong'])
        
        # 设置段落格式
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = self.config.FIRST_LINE_INDENT
        paragraph_format.line_spacing = self.config.LINE_SPACING
        paragraph_format.space_after = Pt(0)
        paragraph_format.space_before = Pt(0)
    
    def _add_heading(self, content: str, level: int):
        """添加标题"""
        paragraph = self.doc.add_paragraph()
        
        run = paragraph.add_run(content)
        
        if level == 1:
            # 一级标题：黑体，三号，不加粗
            run.font.name = self.config.FONTS['heiti']
            run.font.size = self.config.FONT_SIZES['body']
            run.bold = False
            self._set_chinese_font(run, self.config.FONTS['heiti'])
            paragraph.alignment = self.config.ALIGNMENTS['justify']
        elif level == 2:
            # 二级标题：楷体，三号，不加粗
            run.font.name = self.config.FONTS['kaiti']
            run.font.size = self.config.FONT_SIZES['body']
            run.bold = False
            self._set_chinese_font(run, self.config.FONTS['kaiti'])
            paragraph.alignment = self.config.ALIGNMENTS['justify']
        elif level == 3:
            # 三级标题：仿宋，三号，不加粗
            run.font.name = self.config.FONTS['fangsong']
            run.font.size = self.config.FONT_SIZES['body']
            run.bold = False
            self._set_chinese_font(run, self.config.FONTS['fangsong'])
            paragraph.alignment = self.config.ALIGNMENTS['justify']
        else:
            # 其他标题：仿宋，三号，不加粗
            run.font.name = self.config.FONTS['fangsong']
            run.font.size = self.config.FONT_SIZES['body']
            run.bold = False
            self._set_chinese_font(run, self.config.FONTS['fangsong'])
            paragraph.alignment = self.config.ALIGNMENTS['justify']
        
        # 设置段落格式
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = self.config.FIRST_LINE_INDENT  # 标题也要首行缩进2字符
        paragraph_format.line_spacing = self.config.LINE_SPACING
        paragraph_format.space_after = Pt(6)
        paragraph_format.space_before = Pt(6)
    
    def _add_attachment(self, attachment: str):
        """添加附件说明"""
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = self.config.ALIGNMENTS['justify']  # 两端对齐
        
        run = paragraph.add_run(attachment)
        run.font.name = self.config.FONTS['fangsong']
        run.font.size = self.config.FONT_SIZES['body']
        
        self._set_chinese_font(run, self.config.FONTS['fangsong'])
        
        # 设置段落格式
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = self.config.FIRST_LINE_INDENT  # 首行缩进2字符
        paragraph_format.line_spacing = self.config.LINE_SPACING
        paragraph_format.space_after = Pt(0)
    
    def _add_date(self, date: str):
        """添加成文日期"""
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = self.config.ALIGNMENTS['right']
        
        run = paragraph.add_run(date)
        run.font.name = self.config.FONTS['fangsong']
        run.font.size = self.config.FONT_SIZES['body']
        
        self._set_chinese_font(run, self.config.FONTS['fangsong'])
        
        # 设置段落格式
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = self.config.LINE_SPACING
        paragraph_format.space_after = Pt(0)
        paragraph_format.space_before = Pt(12)
    
    def _add_page_numbers(self):
        """添加页码"""
        section = self.doc.sections[0]
        footer = section.footer
        
        # 清除现有内容
        footer.paragraphs[0].clear()
        
        # 添加页码
        paragraph = footer.paragraphs[0]
        paragraph.alignment = self.config.ALIGNMENTS['center']
        
        run = paragraph.add_run("- ")
        run.font.name = self.config.FONTS['fangsong']
        run.font.size = self.config.FONT_SIZES['page_num']
        self._set_chinese_font(run, self.config.FONTS['fangsong'])
        
        # 插入页码字段
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run2 = paragraph.add_run()
        run2._r.append(fldChar1)
        run2._r.append(instrText)
        run2._r.append(fldChar2)
        run2.font.name = self.config.FONTS['fangsong']
        run2.font.size = self.config.FONT_SIZES['page_num']
        self._set_chinese_font(run2, self.config.FONTS['fangsong'])
        
        run3 = paragraph.add_run(" -")
        run3.font.name = self.config.FONTS['fangsong']
        run3.font.size = self.config.FONT_SIZES['page_num']
        self._set_chinese_font(run3, self.config.FONTS['fangsong'])
    
    def _set_chinese_font(self, run, font_name: str):
        """设置中文字体"""
        rPr = run._element.rPr
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            run._element.insert(0, rPr)
        
        # 设置中文字体
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:hint'), 'eastAsia')